"""
IPYNB Analyzer - Gera documentação Word a partir de notebooks Jupyter/Databricks (.ipynb)
Uso: python ipynb_analyzer.py <arquivo.ipynb> [--output saida.docx]
"""

import json
import os
import re
import sys
import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


# ==============================================================================
# CONSTANTES DE CORES (mesma paleta do PBIX Analyzer)
# ==============================================================================

JD_YELLOW     = RGBColor(0xFF, 0xBE, 0x00)
JD_DARK       = RGBColor(0x1A, 0x1A, 0x1A)
DARK_GRAY     = RGBColor(0x40, 0x40, 0x40)
MID_GRAY      = RGBColor(0x70, 0x70, 0x70)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BLUE_CODE     = RGBColor(0x00, 0x50, 0x9D)

GREEN_BG      = "FFBE00"
DARK_GREEN_BG = "1A1A1A"
YELLOW_ACCENT = "FFBE00"
LIGHT_GRAY_BG = "F5F5F5"
ALT_ROW_BG    = "FFF8DC"
CODE_BG       = "F0F4F8"
SQL_BG        = "EAF4EA"
MD_BG         = "FAFAFA"

# ── Dracula Theme ─────────────────────────────────────────────────────────────
DRACULA_BG      = "383A59"   # levemente mais claro que o original #282A36
DRACULA_FG      = RGBColor(0xF8, 0xF8, 0xF2)
DRACULA_COMMENT = RGBColor(0x62, 0x72, 0xA4)
DRACULA_CYAN    = RGBColor(0x8B, 0xE9, 0xFD)
DRACULA_GREEN   = RGBColor(0x50, 0xFA, 0x7B)
DRACULA_ORANGE  = RGBColor(0xFF, 0xB8, 0x6C)
DRACULA_PINK    = RGBColor(0xFF, 0x79, 0xC6)
DRACULA_PURPLE  = RGBColor(0xBD, 0x93, 0xF9)
DRACULA_RED     = RGBColor(0xFF, 0x55, 0x55)
DRACULA_YELLOW  = RGBColor(0xF1, 0xFA, 0x8C)


def _dracula_color(ttype):
    """Mapeia tipo de token pygments para cor Dracula."""
    try:
        from pygments import token as T
        if ttype in T.Comment:              return DRACULA_COMMENT
        if ttype in T.Literal.String:      return DRACULA_YELLOW
        if ttype in T.Keyword:             return DRACULA_PINK
        if ttype in T.Name.Decorator:      return DRACULA_GREEN
        if ttype in T.Name.Function:       return DRACULA_GREEN
        if ttype in T.Name.Class:          return DRACULA_CYAN
        if ttype in T.Name.Builtin:        return DRACULA_CYAN
        if ttype in T.Name.Exception:      return DRACULA_RED
        if ttype in T.Literal.Number:      return DRACULA_PURPLE
        if ttype in T.Operator:            return DRACULA_PINK
        if ttype in T.Punctuation:         return DRACULA_FG
        if ttype in T.Name.Namespace:      return DRACULA_CYAN
        if ttype in T.Keyword.Type:        return DRACULA_CYAN
    except Exception:
        pass
    return DRACULA_FG

# ==============================================================================
# MAPEAMENTO DE MAGIC COMMANDS
# ==============================================================================

MAGIC_LABELS = {
    "%sql":    "SQL",
    "%python": "Python",
    "%scala":  "Scala",
    "%r":      "R",
    "%md":     "Markdown",
    "%sh":     "Shell",
    "%run":    "Run Notebook",
    "%fs":     "DBFS (FileSystem)",
    "%pip":    "Pip Install",
}

# Padrões de fontes de dados
DATA_SOURCE_PATTERNS = [
    (r'spark\.read\.[a-zA-Z]+\s*\(', "spark.read"),
    (r'spark\.read\.table\s*\(\s*["\']([^"\']+)["\']', "spark.read.table"),
    (r'spark\.sql\s*\(\s*["\']([^"\']+)["\']', "spark.sql"),
    (r'dbutils\.fs\.[a-zA-Z]+\s*\(', "dbutils.fs"),
    (r'pd\.read_csv\s*\(', "pd.read_csv"),
    (r'pd\.read_excel\s*\(', "pd.read_excel"),
    (r'pd\.read_parquet\s*\(', "pd.read_parquet"),
    (r'pd\.read_json\s*\(', "pd.read_json"),
    (r'\.load\s*\(\s*["\']([^"\']+)["\']', "load()"),
    (r'FROM\s+([a-zA-Z_][a-zA-Z0-9_.]+)', "SQL FROM"),
    (r'JOIN\s+([a-zA-Z_][a-zA-Z0-9_.]+)', "SQL JOIN"),
]

# Operações DataFrame — (padrão regex, label limpo)
DATAFRAME_OPS = [
    (r'\.groupBy\s*\(',       'groupBy'),
    (r'\.agg\s*\(',           'agg'),
    (r'\.join\s*\(',          'join'),
    (r'\.filter\s*\(',        'filter'),
    (r'\.where\s*\(',         'where'),
    (r'\.select\s*\(',        'select'),
    (r'\.withColumn\s*\(',    'withColumn'),
    (r'\.drop\s*\(',          'drop'),
    (r'\.dropDuplicates\s*\(','dropDuplicates'),
    (r'\.orderBy\s*\(',       'orderBy'),
    (r'\.sort\s*\(',          'sort'),
    (r'\.limit\s*\(',         'limit'),
    (r'\.union\s*\(',         'union'),
    (r'\.pivot\s*\(',         'pivot'),
    (r'\.cache\s*\(',         'cache'),
    (r'\.persist\s*\(',       'persist'),
    (r'\.write\.',            'write'),
    (r'\.saveAsTable\s*\(',   'saveAsTable'),
    (r'\.createOrReplaceTempView\s*\(', 'createOrReplaceTempView'),
    (r'\.toPandas\s*\(',      'toPandas'),
    (r'\.merge\s*\(',         'merge'),
    (r'\.fillna\s*\(',        'fillna'),
    (r'\.dropna\s*\(',        'dropna'),
    (r'\.rename\s*\(',        'rename'),
]


# ==============================================================================
# CAMADA 1: PARSING
# ==============================================================================

def parse_notebook(ipynb_path: str) -> dict:
    """Lê um arquivo .ipynb e retorna estrutura normalizada."""
    with open(ipynb_path, "r", encoding="utf-8-sig") as f:
        content = f.read().strip()

    if not content:
        raise ValueError(
            f"O arquivo '{Path(ipynb_path).name}' está vazio ou corrompido.\n"
            "Re-exporte o notebook do Databricks: File → Export → IPython Notebook (.ipynb)"
        )

    try:
        raw = json.loads(content)
    except json.JSONDecodeError as e:
        raise ValueError(
            f"O arquivo não é um JSON válido: {e}\n"
            "Verifique se o arquivo foi exportado corretamente do Databricks."
        )

    nb_meta   = raw.get("metadata", {})
    kernelspec = nb_meta.get("kernelspec", {})
    lang_info  = nb_meta.get("language_info", {})

    notebook = {
        "path":     ipynb_path,
        "name":     Path(ipynb_path).stem,
        "nbformat": raw.get("nbformat", 4),
        "kernel":   kernelspec.get("display_name", "Desconhecido"),
        "language": lang_info.get("name", kernelspec.get("language", "python")),
        "cells":    [],
    }

    for idx, cell in enumerate(raw.get("cells", [])):
        notebook["cells"].append(_parse_cell(cell, idx))

    return notebook


def _parse_cell(cell: dict, idx: int) -> dict:
    """Normaliza uma célula do notebook."""
    cell_type = cell.get("cell_type", "code")
    source    = cell.get("source", [])

    # source pode ser lista de strings ou string única
    if isinstance(source, list):
        source_str = "".join(source)
    else:
        source_str = source

    outputs     = cell.get("outputs", [])
    exec_count  = cell.get("execution_count")

    magic_cmd, lang, clean_source = _detect_magic(source_str, cell_type)

    # Extrai output textual
    output_text = _extract_output_text(outputs)

    return {
        "index":       idx,
        "type":        cell_type,       # "code", "markdown", "raw"
        "source":      source_str,
        "clean_source": clean_source,   # source sem prefixos MAGIC
        "magic":       magic_cmd,       # "%sql", "%python", None, etc.
        "language":    lang,            # linguagem efetiva da célula
        "outputs":     outputs,
        "output_text": output_text,
        "exec_count":  exec_count,
        "line_count":  len(source_str.splitlines()),
    }


def _detect_magic(source: str, cell_type: str) -> tuple:
    """
    Detecta magic command Databricks numa célula.
    Retorna (magic_cmd, language, clean_source).

    Suporta dois formatos Databricks:
      - Formato .ipynb nativo:  primeira linha = "%sql"
      - Formato exportado:      linhas prefixadas com "# MAGIC %sql"
    """
    if cell_type == "markdown":
        return (None, "markdown", source)
    if cell_type == "raw":
        return (None, "raw", source)

    lines = source.splitlines()
    if not lines:
        return (None, "python", source)

    first = lines[0].strip()

    # Formato "# MAGIC %sql" (Databricks export)
    if first.startswith("# MAGIC"):
        magic_line = re.sub(r'^#\s*MAGIC\s*', '', first).strip()
        for key in MAGIC_LABELS:
            if magic_line.lower().startswith(key):
                # Remove prefixos "# MAGIC" de todas as linhas
                clean_lines = []
                for ln in lines[1:]:
                    clean_ln = re.sub(r'^#\s*MAGIC\s*', '', ln)
                    clean_lines.append(clean_ln)
                return (key, MAGIC_LABELS[key], "\n".join(clean_lines))
        return (None, "python", source)

    # Formato direto: primeira linha é o magic
    for key in MAGIC_LABELS:
        if first.lower() == key or first.lower().startswith(key + " "):
            clean = "\n".join(lines[1:]).strip()
            return (key, MAGIC_LABELS[key], clean)

    # Magic desconhecido (ex: %skip, %restart_python, %conda) — trata como Python
    if first.startswith("%"):
        return (first.split()[0], first.split()[0].lstrip("%").capitalize(), "\n".join(lines[1:]).strip())

    return (None, "python", source)


def _extract_output_text(outputs: list) -> str:
    """Extrai texto legível dos outputs de uma célula."""
    parts = []
    for out in outputs:
        out_type = out.get("output_type", "")
        if out_type in ("stream",):
            text = out.get("text", [])
            if isinstance(text, list):
                text = "".join(text)
            parts.append(text.strip()[:500])
        elif out_type in ("execute_result", "display_data"):
            data = out.get("data", {})
            plain = data.get("text/plain", [])
            if isinstance(plain, list):
                plain = "".join(plain)
            if plain:
                parts.append(plain.strip()[:500])
        elif out_type == "error":
            ename = out.get("ename", "")
            parts.append(f"❌ {ename}: {out.get('evalue', '')}"[:200])
    return "\n".join(parts)


# ==============================================================================
# CAMADA 2: ANÁLISE / INTELIGÊNCIA
# ==============================================================================

def analyze_notebook(notebook: dict) -> dict:
    """Analisa o notebook e extrai inteligência sobre seu conteúdo."""
    cells    = notebook["cells"]
    analysis = {
        "total_cells":     len(cells),
        "code_cells":      0,
        "markdown_cells":  0,
        "sql_cells":       0,
        "languages_used":  set(),
        "imports":         [],
        "data_sources":    [],
        "sql_queries":     [],
        "df_operations":   set(),
        "functions":       [],
        "classes":         [],
        "widgets":         [],
        "temp_views":      [],
        "write_ops":       [],
        "cell_summaries":  [],
        "errors_found":    0,
        "has_outputs":     0,
    }

    seen_imports = set()
    seen_sources = set()
    seen_queries = set()

    for cell in cells:
        ctype = cell["type"]
        src   = cell["source"]
        clean = cell["clean_source"]
        magic = cell["magic"]
        lang  = cell["language"]

        try:  # captura localização exata de qualquer erro de regex
            # Contadores
            if ctype == "code":
                analysis["code_cells"] += 1
            elif ctype == "markdown":
                analysis["markdown_cells"] += 1

            if cell["output_text"]:
                analysis["has_outputs"] += 1
            if "❌" in cell["output_text"]:
                analysis["errors_found"] += 1

            # Linguagens usadas
            if lang and lang not in ("raw", "markdown"):
                analysis["languages_used"].add(lang)

            # Células SQL
            if magic == "%sql" or lang == "SQL":
                analysis["sql_cells"] += 1
                q = clean.strip()
                if q and q not in seen_queries:
                    seen_queries.add(q)
                    analysis["sql_queries"].append({
                        "cell_index": cell["index"],
                        "query":      q,
                        "tables":     _extract_sql_tables(q),
                    })

            if ctype == "code":
                source_to_scan = src  # scan no source original para capturar tudo

                # Imports
                for imp in _extract_imports(source_to_scan):
                    if imp not in seen_imports:
                        seen_imports.add(imp)
                        analysis["imports"].append(imp)

                # Fontes de dados
                for ds in _extract_data_sources(source_to_scan):
                    key = ds["type"] + ds.get("detail", "")
                    if key not in seen_sources:
                        seen_sources.add(key)
                        ds["cell_index"] = cell["index"]
                        analysis["data_sources"].append(ds)

                # spark.sql() embutido no código Python
                for m in re.finditer(r'spark\.sql\s*\(\s*(?:f?["\']|f?""")(.*?)(?:["\']|""")\s*\)',
                                     source_to_scan, re.DOTALL):
                    q = m.group(1).strip()
                    if q and q not in seen_queries:
                        seen_queries.add(q)
                        analysis["sql_queries"].append({
                            "cell_index": cell["index"],
                            "query":      q,
                            "tables":     _extract_sql_tables(q),
                        })

                # Operações DataFrame
                for pat, label in DATAFRAME_OPS:
                    if re.search(pat, source_to_scan):
                        analysis["df_operations"].add(label)

                # Funções definidas
                for m in re.finditer(r'^def\s+(\w+)\s*\(([^)]*)\)', source_to_scan, re.MULTILINE):
                    analysis["functions"].append({
                        "name":   m.group(1),
                        "params": m.group(2).strip(),
                        "cell":   cell["index"],
                    })

                # Classes definidas
                for m in re.finditer(r'^class\s+(\w+)', source_to_scan, re.MULTILINE):
                    analysis["classes"].append({
                        "name": m.group(1),
                        "cell": cell["index"],
                    })

                # Widgets Databricks
                for m in re.finditer(r'dbutils\.widgets\.\w+\s*\(\s*["\']([^"\']+)["\']',
                                     source_to_scan):
                    analysis["widgets"].append(m.group(1))

                # Temp Views
                for m in re.finditer(r'createOrReplaceTempView\s*\(\s*["\']([^"\']+)["\']',
                                     source_to_scan):
                    analysis["temp_views"].append({
                        "name": m.group(1),
                        "cell": cell["index"],
                    })

                # Operações de escrita
                for pat in [r'\.write\.', r'saveAsTable', r'to_csv', r'to_parquet',
                            r'to_excel', r'\.save\s*\(']:
                    if re.search(pat, source_to_scan):
                        analysis["write_ops"].append({
                            "cell":    cell["index"],
                            "pattern": pat.strip("\\.()")
                        })
                        break

            # Resumo da célula
            analysis["cell_summaries"].append(_summarize_cell(cell))

        except Exception as _cell_err:
            import traceback as _tb
            _detail = _tb.format_exc()
            raise RuntimeError(
                f"Erro na célula #{cell['index']+1} (magic={magic!r}, lang={lang!r}):\n"
                f"{_cell_err}\n\nTraceback completo:\n{_detail}"
            ) from _cell_err

    # Finalize
    analysis["languages_used"] = sorted(analysis["languages_used"])
    analysis["df_operations"]  = sorted(analysis["df_operations"])

    # Infere propósito geral do notebook
    analysis["purpose"] = _infer_purpose(analysis)

    return analysis


def _extract_imports(source: str) -> list:
    """Extrai declarações de import do código Python."""
    imports = []
    for line in source.splitlines():
        line = line.strip()
        if line.startswith("import ") or line.startswith("from "):
            # Normaliza
            clean = re.sub(r'\s+', ' ', line)
            imports.append(clean)
    return imports


def _extract_data_sources(source: str) -> list:
    """Extrai fontes de dados do código."""
    found = []
    for pattern, ds_type in DATA_SOURCE_PATTERNS:
        for m in re.finditer(pattern, source, re.IGNORECASE):
            detail = m.group(1) if m.lastindex else ""
            found.append({"type": ds_type, "detail": detail.strip()})
    return found


def _extract_sql_tables(query: str) -> list:
    """Extrai nomes de tabelas de uma query SQL."""
    tables = []
    for kw in ["FROM", "JOIN", "INTO", "UPDATE", "TABLE"]:
        for m in re.finditer(
                rf'\b{kw}\b\s+([`"]?[a-zA-Z_][a-zA-Z0-9_.`"]*)',
                query, re.IGNORECASE):
            t = m.group(1).strip('`"').strip()
            if t.upper() not in ("SELECT", "WHERE", "ON", "SET", "VALUES") and t:
                tables.append(t)
    return list(dict.fromkeys(tables))  # dedup preservando ordem


def _summarize_cell(cell: dict) -> dict:
    """Gera um resumo legível de uma célula."""
    ctype  = cell["type"]
    magic  = cell["magic"]
    lang   = cell["language"]
    src    = cell["clean_source"]
    lines  = cell["line_count"]

    # Título da célula: pega primeiro comentário ou linha de markdown
    title = ""
    for line in src.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip()
            break
        if stripped:
            title = stripped[:80]
            break

    # Tipo legível
    if ctype == "markdown":
        type_label = "📝 Markdown"
    elif magic == "%sql":
        type_label = "🗄️  SQL"
    elif magic in ("%sh", "%fs"):
        type_label = "💻 Shell/FS"
    elif magic == "%run":
        type_label = "▶  Run Notebook"
    elif magic == "%pip":
        type_label = "📦 Pip Install"
    elif ctype == "code":
        type_label = "🐍 Python"
    else:
        type_label = ctype.capitalize()

    # Finalidade inferida
    purpose = _infer_cell_purpose(cell)

    return {
        "index":       cell["index"] + 1,
        "type_label":  type_label,
        "title":       title or "(sem título)",
        "lines":       lines,
        "purpose":     purpose,
        "description": _describe_cell_heuristic(cell),
        "has_output":  bool(cell["output_text"]),
        "exec_count":  cell["exec_count"],
    }


def _infer_cell_purpose(cell: dict) -> str:
    """Infere a finalidade de uma célula pelo seu conteúdo."""
    src   = cell["source"].lower()
    magic = cell["magic"]

    if cell["type"] == "markdown":
        return "Documentação / Seção"
    if magic == "%pip":
        return "Instalação de biblioteca"
    if magic == "%run":
        return "Execução de notebook externo"
    if magic in ("%sh", "%fs"):
        return "Operação de sistema de arquivos"
    if magic == "%sql":
        if any(k in src for k in ["create table", "create or replace"]):
            return "Criação de tabela"
        if "insert" in src:
            return "Inserção de dados"
        if "merge" in src:
            return "Merge / Upsert"
        return "Consulta SQL"
    if "import " in src or src.strip().startswith("from "):
        return "Importação de bibliotecas"
    if "spark.read" in src or "pd.read" in src or "dbutils.fs" in src:
        return "Leitura de dados"
    if ".write." in src or "saveastable" in src or "to_csv" in src:
        return "Escrita / Persistência"
    if "def " in src or "class " in src:
        return "Definição de função/classe"
    if "groupby" in src or "agg(" in src or ".agg(" in src:
        return "Agregação / Sumarização"
    if "join" in src:
        return "Junção de dados"
    if "display(" in src or "show(" in src or "print(" in src:
        return "Visualização / Output"
    if "dbutils.widgets" in src:
        return "Parâmetros (Widgets)"
    if "createorreplacetempview" in src:
        return "Criação de View Temporária"
    return "Processamento / Transformação"


def _describe_cell_heuristic(cell: dict) -> str:
    """Gera descrição contextual de 1-3 frases sobre o que a célula faz."""
    src   = cell["source"]
    clean = cell["clean_source"]
    magic = cell["magic"]
    ctype = cell["type"]
    src_lower = src.lower()

    if ctype == "markdown":
        lines = [l.strip() for l in clean.splitlines()
                 if l.strip() and not l.strip().startswith("#")]
        text = " ".join(lines)[:400]
        return text or "Célula de documentação / seção do notebook."

    if magic == "%pip":
        pkgs = re.findall(r'install\s+([^\n\r;\\]+)', src)
        if pkgs:
            return f"Instala as dependências necessárias: {pkgs[0].strip()[:120]}."
        return "Instala dependências via pip antes da execução do notebook."

    if magic == "%run":
        m = re.search(r'%run\s+(.+)', src)
        if m:
            return f"Executa o notebook externo '{m.group(1).strip()[:100]}', reutilizando seu código ou resultados."
        return "Executa um notebook externo, incorporando seu fluxo ao pipeline atual."

    if magic in ("%sh", "%fs"):
        lines = [l.strip() for l in clean.splitlines() if l.strip()]
        op = lines[0][:120] if lines else ""
        return f"Operação de sistema de arquivos / shell: {op}." if op else "Operação de sistema de arquivos Databricks."

    if magic == "%sql":
        sql_up = clean.upper()
        tables = _extract_sql_tables(clean)
        tlist = ", ".join(tables[:4]) if tables else None

        if "CREATE TABLE" in sql_up or "CREATE OR REPLACE TABLE" in sql_up:
            return f"Cria {'ou recria ' if 'OR REPLACE' in sql_up else ''}a tabela {tlist or 'de destino'}, definindo sua estrutura e dados iniciais."
        if "DROP TABLE" in sql_up:
            m = re.search(r'DROP\s+TABLE(?:\s+IF\s+EXISTS)?\s+([^\s;,()\n]+)', clean, re.IGNORECASE)
            tname = m.group(1) if m else (tlist or "especificada")
            return f"Remove a tabela {tname} do catálogo (geralmente parte de um reset ou recriação)."
        if "MERGE" in sql_up:
            return f"Realiza operação de merge/upsert na tabela {tlist or 'de destino'}, inserindo novos registros e atualizando existentes."
        if "INSERT" in sql_up:
            return f"Insere dados na tabela {tlist or 'de destino'}."
        if "UPDATE" in sql_up:
            return f"Atualiza registros da tabela {tlist or 'de destino'} com base em condição."
        if "DELETE" in sql_up:
            return f"Remove registros da tabela {tlist or 'de destino'} com base em condição."
        if "CREATE" in sql_up and "VIEW" in sql_up:
            return f"Cria uma view SQL sobre as tabelas {tlist or 'base'}."
        if tlist:
            return f"Consulta e transforma dados das tabelas: {tlist}. Retorna o resultado para análise ou uso nas próximas etapas."
        return "Executa consulta SQL para análise ou transformação de dados."

    # ── Python ────────────────────────────────────────────────────────────────
    parts = []

    # Importações
    import_lines = [l.strip() for l in src.splitlines()
                    if l.strip().startswith("import ") or l.strip().startswith("from ")]
    non_import = [l for l in src.strip().splitlines()
                  if l.strip() and not l.strip().startswith("#")
                  and not l.strip().startswith("import ")
                  and not l.strip().startswith("from ")]
    if import_lines and not non_import:
        libs = []
        for imp in import_lines[:6]:
            m = re.match(r'(?:import|from)\s+([\w]+)', imp)
            if m:
                libs.append(m.group(1))
        libs = list(dict.fromkeys(libs))
        return f"Importa as bibliotecas necessárias: {', '.join(libs)}." if libs else "Importa bibliotecas e módulos."

    # Leitura de dados
    read_ops = []
    for pattern, label in [
        (r'spark\.read\.\w+\s*\(', "Spark"),
        (r'spark\.table\s*\(', "Spark (tabela)"),
        (r'pd\.read_csv', "CSV"),
        (r'pd\.read_excel', "Excel"),
        (r'pd\.read_parquet', "Parquet"),
        (r'pd\.read_json', "JSON"),
        (r'dbutils\.fs\.(ls|cp|mv|rm)', "DBFS"),
    ]:
        if re.search(pattern, src):
            read_ops.append(label)

    # Tabelas lidas via spark.read.table / spark.table
    table_reads = re.findall(r'(?:spark\.read\.table|spark\.table)\s*\(\s*["\']([^"\']+)["\']', src, re.IGNORECASE)

    # Transformações
    transforms = []
    for kw, label in [
        (r'\.groupBy\s*\(|\.groupby\s*\(',  "agrupamento"),
        (r'\.join\s*\(',                     "junção de tabelas"),
        (r'\.filter\s*\(|\.where\s*\(',      "filtragem"),
        (r'\.withColumn\s*\(',               "colunas calculadas"),
        (r'\.select\s*\(',                   "seleção de colunas"),
        (r'\.agg\s*\(',                      "agregação"),
        (r'\.merge\s*\(',                    "merge"),
        (r'\.fillna\s*\(|\.dropna\s*\(',     "tratamento de nulos"),
        (r'\.dropDuplicates\s*\(',           "remoção de duplicatas"),
        (r'\.union\s*\(',                    "união de DataFrames"),
        (r'\.pivot\s*\(',                    "pivot"),
    ]:
        if re.search(kw, src):
            transforms.append(label)

    # Escrita
    write_ops = []
    for kw, label in [
        (r'\.saveAsTable\s*\(',  "salva como tabela Delta"),
        (r'\.write\.',           "persiste os dados"),
        (r'to_csv\s*\(',         "exporta para CSV"),
        (r'to_parquet\s*\(',     "exporta para Parquet"),
        (r'to_excel\s*\(',       "exporta para Excel"),
    ]:
        if re.search(kw, src):
            write_ops.append(label)

    # Funções/classes definidas
    func_names = re.findall(r'^def\s+(\w+)', src, re.MULTILINE)
    class_names = re.findall(r'^class\s+(\w+)', src, re.MULTILINE)

    # Widgets
    widgets = re.findall(r'dbutils\.widgets\.\w+\s*\(\s*["\']([^"\']+)["\']', src)

    # Monta descrição
    if func_names or class_names:
        defs = [f"'{n}'" for n in (func_names + class_names)[:4]]
        s = f"Define {'a função' if len(func_names)==1 and not class_names else 'as funções/classes'} {', '.join(defs)}."
        if read_ops or transforms:
            s += f" Também {'lê dados e ' if read_ops else ''}aplica {', '.join(transforms[:2])}." if transforms else ""
        return s

    if widgets:
        return f"Define os parâmetros de entrada do notebook via widgets: {', '.join(widgets[:4])}."

    if read_ops and table_reads and transforms and write_ops:
        return (f"Lê os dados das tabelas {', '.join(table_reads[:3])} via {', '.join(read_ops)}, "
                f"aplica {', '.join(transforms[:3])} e {', '.join(write_ops[:2])}.")

    if read_ops and table_reads and not transforms:
        return f"Carrega os dados das tabelas {', '.join(table_reads[:3])} em DataFrame para as próximas etapas."

    if read_ops and table_reads and transforms:
        return f"Lê as tabelas {', '.join(table_reads[:3])} e aplica {', '.join(transforms[:3])}."

    if read_ops and not table_reads:
        return f"Lê dados via {', '.join(read_ops)} e carrega em DataFrame."

    if transforms and write_ops:
        return f"Transforma os dados ({', '.join(transforms[:3])}) e {', '.join(write_ops[:2])}."

    if transforms:
        return f"Processa e transforma os dados: {', '.join(transforms[:4])}."

    if write_ops:
        return f"Persiste os dados processados: {', '.join(write_ops[:2])}."

    has_display = bool(re.search(r'display\s*\(|\.show\s*\(|print\s*\(', src))
    if has_display:
        return "Exibe os dados ou resultados para visualização e validação."

    # Fallback: primeiro comentário útil
    for line in src.splitlines():
        stripped = line.strip()
        if stripped.startswith("#") and len(stripped) > 5:
            return stripped.lstrip("#").strip()

    return "Executa processamento e transformações nos dados."


def _infer_purpose(analysis: dict) -> str:
    """Infere o propósito geral do notebook a partir da análise."""
    has_sql    = analysis["sql_cells"] > 2
    has_writes = len(analysis["write_ops"]) > 0
    has_reads  = len(analysis["data_sources"]) > 0
    has_funcs  = len(analysis["functions"]) > 0
    langs      = analysis["languages_used"]

    if has_sql and has_writes:
        return "Pipeline de dados (ETL/ELT) — lê, transforma e persiste dados"
    if has_sql and not has_writes:
        return "Análise exploratória — consultas e transformações SQL/Python"
    if has_reads and has_writes and not has_sql:
        return "Processamento Python — lê, transforma e salva com PySpark/Pandas"
    if has_funcs and not has_reads:
        return "Biblioteca de funções / utilitários"
    if "SQL" in langs and "Python" in langs:
        return "Análise híbrida SQL + Python"
    if "SQL" in langs:
        return "Análise e transformação SQL"
    return "Notebook de análise e processamento de dados"


# ==============================================================================
# LLM INTEGRATION
# ==============================================================================

def _enhance_cells_with_llm(notebook: dict, analysis: dict,
                             llm_config: dict, log_callback=None) -> dict:
    """
    Gera guias por célula usando LLM — uma chamada JSON por célula.
    Mesmo padrão do PBIX (enhance_page_with_llm): JSON confiável, focado.
    Retorna dict {cell_index: {"contexto": str, "detalhe": str, "atencao": str}}.
    """
    def _log(msg):
        if log_callback:
            log_callback(msg)

    try:
        from pbix_analyzer import call_llm
    except ImportError:
        return {}

    if not llm_config or llm_config.get("provider") == "Desabilitado":
        return {}

    cells      = notebook["cells"]
    nb_name    = notebook.get("name", "notebook")
    all_results: dict = {}
    MAX_CHARS  = 3000   # máx caracteres de código enviados por célula

    def _lang_label(c):
        magic = c["magic"] or ""
        if c["type"] == "markdown":  return "Markdown"
        if magic == "%sql":          return "SQL"
        if magic == "%pip":          return "Pip Install"
        if magic in ("%sh", "%fs"): return "Shell/FS"
        if magic == "%run":          return "Run Notebook"
        return "Python"

    total = len(cells)
    for c in cells:
        cidx = c["index"]
        lang = _lang_label(c)

        # Células markdown vazias ou muito curtas: pula
        src = c["clean_source"] or c["source"]
        if not src.strip() or (c["type"] == "markdown" and len(src.strip()) < 20):
            continue

        # Trunca código muito longo — mantém início + fim (onde geralmente estão retornos)
        if len(src) > MAX_CHARS:
            half = MAX_CHARS // 2
            code_str = src[:half] + f"\n... ({len(src) - MAX_CHARS} chars omitidos) ...\n" + src[-half:]
        else:
            code_str = src

        _log(f"  🤖 Célula {cidx+1}/{total} ({lang})...")

        prompt = (
            f"Você é um engenheiro de dados sênior especialista em Python, PySpark e SQL.\n"
            f"Documente a célula {cidx+1} do notebook '{nb_name}' para um desenvolvedor júnior que vai mantê-lo.\n\n"
            f"LINGUAGEM: {lang}\n"
            f"CÓDIGO:\n{code_str}\n\n"
            f"Escreva em português técnico. Analise o código real:\n"
            f"- Funções definidas: explique o que cada uma faz, parâmetros e retorno\n"
            f"- SQL: tabelas acessadas, filtros, agregações, resultado\n"
            f"- Transformações: o que os dados eram e o que se tornam\n"
            f"- Listas/dicionários de configuração: para que servem\n\n"
            f'Responda SOMENTE em JSON válido:\n'
            f'{{"contexto": "...", "detalhe": "...", "atencao": "..."}}\n\n'
            f"contexto: narrativa completa sobre o propósito, funções e resultado no pipeline\n"
            f"detalhe: decisões técnicas não óbvias ou nuances de negócio (ou '—')\n"
            f"atencao: riscos reais: DROP/overwrite, configs dev×prod, dependências críticas (ou '—')"
        )

        raw = call_llm(prompt, llm_config)
        if not raw:
            _log(f"    ⚠️  Sem resposta do LLM para célula {cidx+1}.")
            continue

        try:
            text = raw.strip()
            if "```" in text:
                text = re.sub(r"```(?:json)?", "", text).strip().rstrip("`").strip()
            # Extrai JSON mesmo que haja texto antes/depois
            m = re.search(r'\{.*\}', text, re.DOTALL)
            if m:
                text = m.group(0)
            import json as _json
            data = _json.loads(text)
            entry = {
                "contexto": str(data.get("contexto", "")).strip(),
                "detalhe":  str(data.get("detalhe",  "")).strip(),
                "atencao":  str(data.get("atencao",  "")).strip(),
            }
            # Filtra campos vazios ou só traço
            for k in ("detalhe", "atencao"):
                if entry[k] in ("", "—", "-", "–", "N/A", "Nenhum", "Nenhum.", "—."):
                    entry[k] = ""
            if entry["contexto"]:
                all_results[cidx] = entry
                _log(f"    ✅ Célula {cidx+1} documentada.")
            else:
                _log(f"    ⚠️  Célula {cidx+1}: JSON vazio.")
        except Exception as e:
            _log(f"    ⚠️  Célula {cidx+1}: erro ao parsear JSON — {e}")

    return all_results

def enhance_notebook_with_llm(notebook: dict, analysis: dict, llm_config: dict):
    """
    Gera narrativa de IA para o notebook.
    Retorna dict com objetivo, resumo e guia, ou None se LLM não configurado/falhar.
    """
    try:
        from pbix_analyzer import call_llm
    except ImportError:
        return None

    if not llm_config or llm_config.get("provider") == "Desabilitado":
        return None

    # Monta contexto estrutural (sem dados reais)
    langs    = ", ".join(analysis["languages_used"]) or "Python"
    imports  = "\n".join(f"  - {i}" for i in analysis["imports"][:20])
    sources  = "\n".join(f"  - {d['type']}: {d.get('detail','')}" for d in analysis["data_sources"][:15])
    sql_list = "\n".join(f"  - Célula {q['cell_index']+1}: {q['query'][:120]}..." for q in analysis["sql_queries"][:8])
    funcs    = ", ".join(f["name"] for f in analysis["functions"][:10])
    writes   = str(len(analysis["write_ops"]))

    prompt = f"""Você é um documentador técnico especialista em engenharia de dados e Databricks.
Analise o seguinte notebook e escreva documentação em português corporativo.

NOTEBOOK: {notebook['name']}
KERNEL: {notebook['kernel']}
LINGUAGENS: {langs}
TOTAL DE CÉLULAS: {analysis['total_cells']} ({analysis['code_cells']} código, {analysis['markdown_cells']} markdown, {analysis['sql_cells']} SQL)
PROPÓSITO INFERIDO: {analysis['purpose']}

IMPORTS/BIBLIOTECAS:
{imports or '  (nenhum detectado)'}

FONTES DE DADOS:
{sources or '  (nenhuma detectada)'}

QUERIES SQL (amostra):
{sql_list or '  (nenhuma)'}

FUNÇÕES DEFINIDAS: {funcs or 'nenhuma'}
OPERAÇÕES DE ESCRITA: {writes}

Escreva em 3 partes — seja específico, técnico e direto ao ponto:

**OBJETIVO:**
(1 parágrafo — o que esse notebook faz e para quem serve)

**RESUMO DO FLUXO:**
(2-3 parágrafos — como o notebook funciona do início ao fim: o que lê, como transforma, o que produz)

**GUIA DE LEITURA:**
(passos numerados — como navegar pelo notebook para entender rapidamente)"""

    raw = call_llm(prompt, llm_config)
    if not raw:
        return None

    result = {"objetivo": "", "resumo": "", "leitura": ""}
    current = None
    buffer  = []

    for line in raw.splitlines():
        ul = line.upper()
        if "OBJETIVO" in ul:
            current = "objetivo"; buffer = []
        elif "RESUMO" in ul:
            if current: result[current] = "\n".join(buffer).strip()
            current = "resumo"; buffer = []
        elif "GUIA" in ul or "LEITURA" in ul:
            if current: result[current] = "\n".join(buffer).strip()
            current = "leitura"; buffer = []
        elif current:
            buffer.append(line)

    if current and buffer:
        result[current] = "\n".join(buffer).strip()

    return result if any(result.values()) else None


# ==============================================================================
# HELPERS DE FORMATAÇÃO WORD (compartilhados com pbix_analyzer)
# ==============================================================================

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color_hex="FFBE00", width_pt=12, sides=None):
    tc      = cell._tc
    tcPr    = tc.get_or_add_tcPr()
    tcBords = OxmlElement("w:tcBorders")
    for side in (sides or ["top", "left", "bottom", "right"]):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(width_pt))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        tcBords.append(b)
    tcPr.append(tcBords)


def _set_para_border_bottom(para, color_hex="FFBE00", size_pt=6):
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt * 8))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_header_row(table, texts, bg_hex="FFBE00"):
    row = table.rows[0]
    for i, text in enumerate(texts):
        if i >= len(row.cells): break
        cell = row.cells[i]
        cell.text = ""
        _set_cell_bg(cell, bg_hex)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(str(text))
        run.font.bold  = True
        run.font.color.rgb = JD_DARK
        run.font.size  = Pt(9)


def _add_data_row(table, row_idx, values, alternate=False):
    row = table.rows[row_idx]
    bg  = ALT_ROW_BG if alternate else "FFFFFF"
    for i, val in enumerate(values):
        if i >= len(row.cells): break
        cell = row.cells[i]
        cell.text = ""
        _set_cell_bg(cell, bg)
        run = cell.paragraphs[0].add_run(str(val) if val is not None else "—")
        run.font.size = Pt(9)


def _section_header(doc, number, title):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, DARK_GREEN_BG)
    _set_cell_border(cell, color_hex="FFBE00", width_pt=4)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(f"  {number}.  {title.upper()}")
    run.font.bold  = True
    run.font.color.rgb = JD_YELLOW
    run.font.size  = Pt(14)
    doc.add_paragraph()
    return tbl


def _subsection(doc, title, level=2):
    p   = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold  = True
    run.font.size  = Pt(12 if level == 2 else 10)
    run.font.color.rgb = JD_DARK if level == 2 else DARK_GRAY
    _set_para_border_bottom(p, color_hex=GREEN_BG, size_pt=2)
    doc.add_paragraph()
    return p


def _info_box(doc, label, content, bg_color="FFFAE6", border_color="FFBE00"):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, bg_color)
    _set_cell_border(cell, color_hex=border_color, width_pt=16)
    para = cell.paragraphs[0]
    if label:
        r = para.add_run(f"{label}  ")
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = JD_DARK
    if isinstance(content, list):
        for i, item in enumerate(content):
            if i == 0:
                r = para.add_run(str(item))
                r.font.size = Pt(9.5)
            else:
                np_ = cell.add_paragraph()
                r   = np_.add_run(str(item))
                r.font.size = Pt(9.5)
    else:
        r = para.add_run(str(content))
        r.font.size = Pt(9.5)
    doc.add_paragraph()


def _code_block(doc, code: str, lang="python", max_lines=35):
    """Renderiza bloco de código com Dracula theme e syntax highlighting via pygments."""
    lines = code.strip().splitlines()
    truncated = len(lines) > max_lines
    display_lines = lines[:max_lines] if truncated else lines
    display_code  = "\n".join(display_lines)

    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, DRACULA_BG)
    _set_cell_border(cell, color_hex="44475A", width_pt=6)

    # Tenta syntax highlighting com pygments
    highlighted = False
    try:
        from pygments import lex
        from pygments.lexers import PythonLexer, get_lexer_by_name
        lexer = get_lexer_by_name("sql") if lang == "SQL" else PythonLexer()
        tokens = list(lex(display_code, lexer))
        highlighted = True
    except Exception:
        tokens = []

    if highlighted:
        para = cell.paragraphs[0]
        # Define espaçamento mínimo do parágrafo
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after  = Pt(1)

        for ttype, value in tokens:
            color = _dracula_color(ttype)
            # Divide por newline: cada \n cria novo parágrafo
            if "\n" in value:
                parts = value.split("\n")
                for i, part in enumerate(parts):
                    if part:
                        run = para.add_run(part)
                        run.font.name      = "Courier New"
                        run.font.size      = Pt(8)
                        run.font.color.rgb = color
                    if i < len(parts) - 1:
                        para = cell.add_paragraph()
                        para.paragraph_format.space_before = Pt(1)
                        para.paragraph_format.space_after  = Pt(1)
            else:
                run = para.add_run(value)
                run.font.name      = "Courier New"
                run.font.size      = Pt(8)
                run.font.color.rgb = color
    else:
        # Fallback simples sem highlighting
        para = cell.paragraphs[0]
        run  = para.add_run(display_code)
        run.font.name      = "Courier New"
        run.font.size      = Pt(8)
        run.font.color.rgb = DRACULA_FG

    if truncated:
        omit_p   = cell.add_paragraph()
        omit_run = omit_p.add_run(f"  ... ({len(lines) - max_lines} linhas omitidas)")
        omit_run.font.size      = Pt(8)
        omit_run.font.italic    = True
        omit_run.font.color.rgb = DRACULA_COMMENT

    doc.add_paragraph()


def _bullets(doc, items, icon="▸"):
    for item in items:
        p   = doc.add_paragraph()
        run = p.add_run(f"  {icon}  {item}")
        run.font.size = Pt(9.5)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)


def _body(doc, text, size=10, bold=False, color=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def _separator(doc):
    p = doc.add_paragraph()
    _set_para_border_bottom(p, color_hex=GREEN_BG, size_pt=1)
    p.paragraph_format.space_after = Pt(4)


# ==============================================================================
# CAMADA 3: GERAÇÃO DO DOCUMENTO WORD
# ==============================================================================

def generate_notebook_docx(nb_name: str, notebook: dict, analysis: dict,
                            output_path: str, llm_config=None, log_callback=None):
    """Gera o documento Word completo para o notebook."""

    def _log(msg):
        if log_callback:
            log_callback(msg)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)

    # ── CAPA ──────────────────────────────────────────────────────────────────
    _log("📄 Gerando capa...")

    cover_tbl  = doc.add_table(rows=1, cols=1)
    cover_tbl.style = "Table Grid"
    cover_cell = cover_tbl.cell(0, 0)
    _set_cell_bg(cover_cell, DARK_GREEN_BG)
    para = cover_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run("  DOCUMENTAÇÃO DE NOTEBOOK DATABRICKS  ")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = JD_YELLOW

    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(nb_name)
    tr.font.bold = True; tr.font.size = Pt(26); tr.font.color.rgb = JD_DARK

    doc.add_paragraph()

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("Documentação Técnica e Analítica")
    sr.font.size = Pt(16); sr.font.bold = True; sr.font.color.rgb = DARK_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    acc_tbl  = doc.add_table(rows=1, cols=1)
    acc_tbl.style = "Table Grid"
    _set_cell_bg(acc_tbl.cell(0, 0), YELLOW_ACCENT)
    ap = acc_tbl.cell(0, 0).paragraphs[0]
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = ap.add_run("  Gerado automaticamente pelo PBIX Analyzer — Módulo Notebooks  ")
    ar.font.size = Pt(9); ar.font.bold = True; ar.font.color.rgb = JD_DARK

    doc.add_paragraph()

    gen_p = doc.add_paragraph()
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_p.add_run(f"Data de Geração: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    kp = doc.add_paragraph()
    kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kp.add_run(f"Kernel: {notebook['kernel']}  |  Linguagem: {notebook['language']}")

    doc.add_paragraph()
    doc.add_paragraph()

    # Stats de capa
    langs_str = ", ".join(analysis["languages_used"]) or notebook["language"]
    stats_data = [
        (str(analysis["total_cells"]),        "Células"),
        (str(analysis["sql_cells"]),           "Queries SQL"),
        (str(len(analysis["data_sources"])),   "Fontes de Dados"),
        (str(len(analysis["functions"])),      "Funções"),
    ]
    stats_tbl = doc.add_table(rows=1, cols=4)
    stats_tbl.style = "Table Grid"
    for i, (val, lbl) in enumerate(stats_data):
        sc = stats_tbl.cell(0, i)
        _set_cell_bg(sc, GREEN_BG)
        sp = sc.paragraphs[0]
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = sp.add_run(val + "\n")
        vr.font.bold = True; vr.font.size = Pt(20); vr.font.color.rgb = JD_DARK
        lr = sp.add_run(lbl)
        lr.font.size = Pt(9); lr.font.color.rgb = JD_DARK

    doc.add_page_break()

    # ── ÍNDICE ────────────────────────────────────────────────────────────────
    idx_p = doc.add_paragraph()
    idx_r = idx_p.add_run("ÍNDICE")
    idx_r.font.bold = True; idx_r.font.size = Pt(16); idx_r.font.color.rgb = JD_DARK
    _set_para_border_bottom(idx_p, color_hex=GREEN_BG, size_pt=3)
    doc.add_paragraph()

    sections_index = [
        "1. Resumo Executivo",
        "2. Ambiente e Bibliotecas",
        "3. Fontes de Dados",
        "4. Queries SQL",
        "5. Fluxo de Processamento",
        "6. Funções e Classes Definidas",
        "7. Outputs e Resultados",
    ]
    for entry in sections_index:
        p   = doc.add_paragraph()
        run = p.add_run(f"  {entry}")
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ── LLM: NARRATIVA GERAL + DESCRIÇÕES POR CÉLULA ─────────────────────────
    llm_result = None
    cell_descriptions: dict = {}  # {cell_index: str}

    if llm_config and llm_config.get("provider") not in (None, "Desabilitado"):
        _log("🤖 Gerando narrativa LLM para o notebook...")
        llm_result = enhance_notebook_with_llm(notebook, analysis, llm_config)
        if llm_result:
            _log("✅ Narrativa LLM gerada com sucesso.")
        else:
            _log("⚠️  LLM não retornou narrativa — usando heurísticas.")

        _log("🤖 Gerando descrições das células com LLM (lotes de 5)...")
        cell_descriptions = _enhance_cells_with_llm(notebook, analysis, llm_config,
                                                     log_callback=_log)
        if cell_descriptions:
            _log(f"✅ {len(cell_descriptions)} células documentadas pelo LLM.")
        else:
            _log("⚠️  LLM não retornou descrições — usando heurísticas por célula.")

    # ── SEÇÃO 1: RESUMO EXECUTIVO ─────────────────────────────────────────────
    _log("📝 Seção 1 — Resumo Executivo...")
    _section_header(doc, 1, "Resumo Executivo")

    if llm_result and llm_result.get("objetivo"):
        _subsection(doc, "Objetivo")
        _body(doc, llm_result["objetivo"])
        doc.add_paragraph()

        if llm_result.get("resumo"):
            _subsection(doc, "Resumo do Fluxo")
            _body(doc, llm_result["resumo"])
            doc.add_paragraph()

        if llm_result.get("leitura"):
            _subsection(doc, "Guia de Leitura")
            _body(doc, llm_result["leitura"])
            doc.add_paragraph()
    else:
        _info_box(doc, "Propósito:", analysis["purpose"])

        stats_items = [
            f"Total de células: {analysis['total_cells']} ({analysis['code_cells']} código | {analysis['markdown_cells']} markdown | {analysis['sql_cells']} SQL)",
            f"Linguagens utilizadas: {langs_str}",
            f"Fontes de dados identificadas: {len(analysis['data_sources'])}",
            f"Queries SQL extraídas: {len(analysis['sql_queries'])}",
            f"Funções definidas: {len(analysis['functions'])}",
            f"Operações de escrita: {len(analysis['write_ops'])}",
        ]
        if analysis["widgets"]:
            stats_items.append(f"Parâmetros (widgets): {', '.join(analysis['widgets'][:5])}")
        if analysis["errors_found"]:
            stats_items.append(f"⚠️  Células com erro nas saídas: {analysis['errors_found']}")

        _bullets(doc, stats_items)

    doc.add_page_break()

    # ── SEÇÃO 2: AMBIENTE E BIBLIOTECAS ───────────────────────────────────────
    _log("📦 Seção 2 — Ambiente e Bibliotecas...")
    _section_header(doc, 2, "Ambiente e Bibliotecas")

    _body(doc, f"Kernel: {notebook['kernel']}  |  Linguagem: {notebook['language']}", bold=True)
    doc.add_paragraph()

    if analysis["imports"]:
        _subsection(doc, "Imports Detectados")
        tbl = doc.add_table(rows=len(analysis["imports"]) + 1, cols=2)
        tbl.style = "Table Grid"
        _add_header_row(tbl, ["#", "Declaração de Import"])
        for i, imp in enumerate(analysis["imports"], 1):
            _add_data_row(tbl, i, [str(i), imp], alternate=(i % 2 == 0))
        # Larguras
        for row in tbl.rows:
            row.cells[0].width = Inches(0.4)
            row.cells[1].width = Inches(5.6)
        doc.add_paragraph()
    else:
        _body(doc, "Nenhum import explícito detectado nas células de código.", color=MID_GRAY)

    if analysis["widgets"]:
        doc.add_paragraph()
        _subsection(doc, "Parâmetros / Widgets Databricks")
        _bullets(doc, analysis["widgets"])

    doc.add_page_break()

    # ── SEÇÃO 3: FONTES DE DADOS ───────────────────────────────────────────────
    _log("🗄️  Seção 3 — Fontes de Dados...")
    _section_header(doc, 3, "Fontes de Dados")

    if analysis["data_sources"]:
        tbl = doc.add_table(rows=len(analysis["data_sources"]) + 1, cols=3)
        tbl.style = "Table Grid"
        _add_header_row(tbl, ["Tipo", "Detalhe / Tabela", "Célula"])
        for i, ds in enumerate(analysis["data_sources"], 1):
            cell_ref = f"#{ds.get('cell_index', 0) + 1}" if "cell_index" in ds else "—"
            _add_data_row(tbl, i, [ds["type"], ds.get("detail", "—"), cell_ref],
                          alternate=(i % 2 == 0))
        doc.add_paragraph()
    else:
        _body(doc, "Nenhuma fonte de dados explícita detectada.", color=MID_GRAY)

    if analysis["temp_views"]:
        doc.add_paragraph()
        _subsection(doc, "Views Temporárias Criadas")
        rows = [(v["name"], f"#{v['cell'] + 1}") for v in analysis["temp_views"]]
        tbl  = doc.add_table(rows=len(rows) + 1, cols=2)
        tbl.style = "Table Grid"
        _add_header_row(tbl, ["Nome da View", "Célula"])
        for i, (name, cel) in enumerate(rows, 1):
            _add_data_row(tbl, i, [name, cel], alternate=(i % 2 == 0))
        doc.add_paragraph()

    doc.add_page_break()

    # ── SEÇÃO 4: QUERIES SQL ──────────────────────────────────────────────────
    _log("🔍 Seção 4 — Queries SQL...")
    _section_header(doc, 4, "Queries SQL")

    if analysis["sql_queries"]:
        _body(doc, f"{len(analysis['sql_queries'])} queries SQL identificadas no notebook.")
        doc.add_paragraph()

        for qi, q in enumerate(analysis["sql_queries"], 1):
            _subsection(doc, f"Query {qi} — Célula #{q['cell_index'] + 1}", level=3)
            if q["tables"]:
                _body(doc, f"Tabelas envolvidas: {', '.join(q['tables'])}", bold=True)
            _code_block(doc, q["query"], lang="SQL", max_lines=40)
    else:
        _body(doc, "Nenhuma query SQL detectada.", color=MID_GRAY)

    doc.add_page_break()

    # ── SEÇÃO 5: FLUXO DE PROCESSAMENTO ──────────────────────────────────────
    _log("⚙️  Seção 5 — Fluxo de Processamento...")
    _section_header(doc, 5, "Fluxo de Processamento")

    _body(doc, f"Abaixo o detalhamento célula a célula das {analysis['total_cells']} células do notebook.")
    doc.add_paragraph()

    # Tabela-resumo de todas as células
    summaries = analysis["cell_summaries"]
    tbl = doc.add_table(rows=len(summaries) + 1, cols=4)
    tbl.style = "Table Grid"
    _add_header_row(tbl, ["#", "Tipo", "Finalidade Inferida", "Linhas"])
    for i, s in enumerate(summaries, 1):
        _add_data_row(tbl, i,
                      [str(s["index"]), s["type_label"], s["purpose"], str(s["lines"])],
                      alternate=(i % 2 == 0))
    for row in tbl.rows:
        row.cells[0].width = Inches(0.3)
        row.cells[1].width = Inches(1.2)
        row.cells[2].width = Inches(3.7)
        row.cells[3].width = Inches(0.6)
    doc.add_paragraph()

    # Detalhamento das células de código (não-markdown, com conteúdo relevante)
    doc.add_paragraph()
    _subsection(doc, "Detalhamento das Células de Código")

    shown = 0
    for cell in notebook["cells"]:
        if cell["type"] == "markdown":
            continue
        if cell["line_count"] == 0:
            continue

        summary = summaries[cell["index"]]
        magic   = cell["magic"] or ""
        lang    = "SQL" if magic == "%sql" else "python"

        # Cabeçalho: número + tipo + finalidade
        hdr_tbl = doc.add_table(rows=1, cols=2)
        hdr_tbl.style = "Table Grid"
        left  = hdr_tbl.cell(0, 0)
        right = hdr_tbl.cell(0, 1)
        bg    = SQL_BG if lang == "SQL" else CODE_BG
        _set_cell_bg(left, DARK_GREEN_BG)
        _set_cell_bg(right, bg)
        left.width = Inches(0.3)
        left_p = left.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ln = left_p.add_run(str(summary["index"]))
        ln.font.bold = True; ln.font.size = Pt(9); ln.font.color.rgb = JD_YELLOW

        right_p = right.paragraphs[0]
        rn = right_p.add_run(f"{summary['type_label']}  —  {summary['purpose']}")
        rn.font.bold = True; rn.font.size = Pt(9); rn.font.color.rgb = JD_DARK

        # ── Guia do Desenvolvedor (LLM estruturado > heurística) ─────────────
        guide = cell_descriptions.get(cell["index"])  # dict com contexto/detalhe/atencao

        if guide and isinstance(guide, dict):
            # LLM estruturado: 3 caixas distintas
            contexto = guide.get("contexto", "")
            detalhe  = guide.get("detalhe",  "")
            atencao  = guide.get("atencao",  "")

            if contexto:
                _info_box(doc, "📋 Contexto:", contexto,
                          bg_color="EBF4FB", border_color="2E86C1")
            if detalhe:
                _info_box(doc, "🔍 Detalhe técnico:", detalhe,
                          bg_color="F4F6F7", border_color="717D7E")
            if atencao:
                _info_box(doc, "⚠️ Atenção:", atencao,
                          bg_color="FEF9E7", border_color="F39C12")
        else:
            # Fallback heurístico
            desc = summary.get("description", "")
            if desc:
                _info_box(doc, "📋 O que faz:", desc,
                          bg_color="EBF4FB", border_color="2E86C1")

        # Código
        _code_block(doc, cell["clean_source"] or cell["source"], lang=lang, max_lines=35)

        # Output (quando disponível e sem erro)
        if cell["output_text"] and "❌" not in cell["output_text"]:
            out_preview = cell["output_text"][:300]
            _info_box(doc, "▶ Output:", out_preview, bg_color="F5F5F5", border_color="CCCCCC")

        shown += 1

    doc.add_page_break()

    # ── SEÇÃO 6: FUNÇÕES E CLASSES ────────────────────────────────────────────
    _log("🔧 Seção 6 — Funções e Classes...")
    _section_header(doc, 6, "Funções e Classes Definidas")

    if analysis["functions"]:
        _subsection(doc, "Funções")
        rows = [(f["name"], f["params"] or "sem parâmetros", f"#{f['cell'] + 1}")
                for f in analysis["functions"]]
        tbl  = doc.add_table(rows=len(rows) + 1, cols=3)
        tbl.style = "Table Grid"
        _add_header_row(tbl, ["Nome", "Parâmetros", "Célula"])
        for i, r in enumerate(rows, 1):
            _add_data_row(tbl, i, r, alternate=(i % 2 == 0))
        doc.add_paragraph()
    else:
        _body(doc, "Nenhuma função definida no notebook.", color=MID_GRAY)

    if analysis["classes"]:
        doc.add_paragraph()
        _subsection(doc, "Classes")
        _bullets(doc, [f"{c['name']}  (célula #{c['cell'] + 1})" for c in analysis["classes"]])
    doc.add_paragraph()

    if analysis["df_operations"]:
        _subsection(doc, "Operações DataFrame Utilizadas")
        _body(doc, "Transformações PySpark/Pandas detectadas:")
        _bullets(doc, analysis["df_operations"])

    doc.add_page_break()

    # ── SEÇÃO 7: OUTPUTS E RESULTADOS ─────────────────────────────────────────
    _log("📊 Seção 7 — Outputs e Resultados...")
    _section_header(doc, 7, "Outputs e Resultados")

    cells_with_output = [c for c in notebook["cells"]
                         if c["output_text"] and "❌" not in c["output_text"]]
    cells_with_errors = [c for c in notebook["cells"] if "❌" in c["output_text"]]

    if cells_with_output:
        _body(doc, f"{len(cells_with_output)} células com outputs registrados.")
        doc.add_paragraph()

        for cell in cells_with_output[:15]:  # limita a 15 outputs
            summary = summaries[cell["index"]]
            _subsection(doc, f"Célula #{summary['index']} — {summary['purpose']}", level=3)
            preview = cell["output_text"][:600]
            _info_box(doc, "Output:", preview, bg_color="F5F5F5", border_color="CCCCCC")
    else:
        _body(doc, "Nenhum output registrado no notebook.", color=MID_GRAY)

    if cells_with_errors:
        doc.add_paragraph()
        _subsection(doc, "⚠️  Células com Erros")
        for cell in cells_with_errors:
            summary = summaries[cell["index"]]
            _info_box(doc, f"Célula #{summary['index']}:",
                      cell["output_text"][:300],
                      bg_color="FFF0F0", border_color="CC0000")

    if analysis["write_ops"]:
        doc.add_paragraph()
        _subsection(doc, "Operações de Escrita / Persistência")
        rows = [(f"#{w['cell'] + 1}", w["pattern"]) for w in analysis["write_ops"]]
        tbl  = doc.add_table(rows=len(rows) + 1, cols=2)
        tbl.style = "Table Grid"
        _add_header_row(tbl, ["Célula", "Operação Detectada"])
        for i, r in enumerate(rows, 1):
            _add_data_row(tbl, i, r, alternate=(i % 2 == 0))
        doc.add_paragraph()

    # ── RODAPÉ ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    footer_tbl = doc.add_table(rows=1, cols=1)
    footer_tbl.style = "Table Grid"
    _set_cell_bg(footer_tbl.cell(0, 0), DARK_GREEN_BG)
    fp = footer_tbl.cell(0, 0).paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(f"Gerado pelo PBIX Analyzer — Módulo Notebooks  |  {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    fr.font.size = Pt(8); fr.font.color.rgb = JD_YELLOW

    doc.save(output_path)
    _log(f"✅ Documento salvo em: {output_path}")


# ==============================================================================
# ENTRY POINT (linha de comando)
# ==============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="IPYNB Analyzer — documentação de notebooks Jupyter/Databricks"
    )
    parser.add_argument("ipynb", help="Caminho para o arquivo .ipynb")
    parser.add_argument("--output", "-o", help="Caminho de saída para o .docx (opcional)")
    args = parser.parse_args()

    ipynb_path = args.ipynb
    if not os.path.exists(ipynb_path):
        print(f"ERRO: Arquivo não encontrado: {ipynb_path}")
        sys.exit(1)

    nb_name = Path(ipynb_path).stem
    if args.output:
        output_path = args.output
    else:
        output_path = str(Path(ipynb_path).parent / f"{nb_name}_Documentacao.docx")

    print(f"📂 Lendo: {ipynb_path}")
    print("📋 Parseando notebook...")
    notebook = parse_notebook(ipynb_path)

    print("🔍 Analisando estrutura...")
    analysis = analyze_notebook(notebook)

    print("📝 Gerando documentação Word...")
    generate_notebook_docx(nb_name, notebook, analysis, output_path)

    print(f"\n✅ Documentação gerada com sucesso!")
    print(f"   📄 Arquivo: {output_path}")
    print(f"   📓 {analysis['total_cells']} células analisadas")
    print(f"   🗄️  {len(analysis['sql_queries'])} queries SQL extraídas")
    print(f"   📦 {len(analysis['imports'])} imports detectados")
    print(f"   🔗 {len(analysis['data_sources'])} fontes de dados identificadas")


if __name__ == "__main__":
    main()
