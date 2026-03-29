"""
PBIX Analyzer - Gera documentação Word rica a partir de arquivos Power BI (.pbix)
Uso: python pbix_analyzer.py <arquivo.pbix> [--output saida.docx]
Arquitetura: Parsing -> Inteligência/Análise -> Renderização de Documento
"""

import argparse
import json
import os
import re
import shutil
import sys
import tempfile
import zipfile
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


# ==============================================================================
# CONSTANTES DE CORES E TIPOGRAFIA
# ==============================================================================

JD_GREEN    = RGBColor(0x1A, 0x1A, 0x1A)   # preto  – texto em fundo claro/branco
JD_YELLOW   = RGBColor(0xFF, 0xBE, 0x00)   # amarelo Construction – texto em fundo escuro
JD_DARK     = RGBColor(0x1A, 0x1A, 0x1A)   # preto
DARK_GRAY   = RGBColor(0x40, 0x40, 0x40)
MID_GRAY    = RGBColor(0x70, 0x70, 0x70)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GOLD        = RGBColor(0xB8, 0x86, 0x00)

GREEN_BG      = "FFBE00"   # amarelo Construction – fundo de cabeçalhos
DARK_GREEN_BG = "1A1A1A"   # preto – fundo escuro (capa, seções)
YELLOW_ACCENT = "FFBE00"   # amarelo Construction – fundo de destaque
LIGHT_GRAY_BG = "F5F5F5"
ALT_ROW_BG    = "FFF8DC"   # bege claro – linhas alternadas (evita verde)

VISUAL_TYPE_LABELS = {
    "card":                             "Cartão (Card)",
    "multiRowCard":                     "Cartão de Múltiplas Linhas",
    "slicer":                           "Segmentação de Dados (Slicer)",
    "advancedSlicerVisual":             "Segmentação Avançada",
    "ChicletSlicer1448559807354":       "Chiclet Slicer (Custom)",
    "tableEx":                          "Tabela",
    "pivotTable":                       "Matriz",
    "lineChart":                        "Gráfico de Linhas",
    "clusteredBarChart":                "Gráfico de Barras Agrupadas",
    "clusteredColumnChart":             "Gráfico de Colunas Agrupadas",
    "columnChart":                      "Gráfico de Colunas",
    "lineClusteredColumnComboChart":    "Gráfico Combinado (Linha + Coluna)",
    "lineStackedColumnComboChart":      "Gráfico Combinado (Linha + Coluna Empilhada)",
    "donutChart":                       "Gráfico de Rosca",
    "pieChart":                         "Gráfico de Pizza",
    "scatterChart":                     "Gráfico de Dispersão",
    "EnhancedScatterChart1443994985041": "Dispersão Aprimorada (Custom)",
    "BoxandWhiskerByMAQ1823AD39DT234AB532063E128AX": "Box e Whisker (Custom)",
    "azureMap":                         "Mapa Azure",
    "map":                              "Mapa",
    "shape":                            "Forma / Elemento Visual",
    "textbox":                          "Caixa de Texto",
    "image":                            "Imagem",
    "actionButton":                     "Botão de Ação",
    "deneb7E15AEF80B9E4D4F8E12924291ECE89A": "Deneb (Custom Vega/Vega-Lite)",
    "STANDALONEdeneb7E15AEF80B9E4D4F8E12924291ECE89A": "Deneb Standalone (Custom)",
    "scatterChartByAkvelon6CFB0DAB29E746BE850B6DFD1E2789FD": "Dispersão Aprimorada (Akvelon)",
    "advancedtoggleswitch": "Toggle Switch Avançado",
}

DECORATION_TYPES = {"shape", "textbox", "image", "actionButton"}
SLICER_TYPES     = {"slicer", "advancedSlicerVisual", "ChicletSlicer1448559807354"}

CUSTOM_VISUAL_MAP = {
    "advancedtoggleswitch": {
        "name": "Advanced Toggle Switch",
        "desc": "Botão de alternância visual que permite ao usuário ativar/desativar filtros ou modos de exibição de forma intuitiva.",
        "value": "Melhora a usabilidade ao substituir slicers convencionais por controles mais visuais e modernos."
    },
    "BoxandWhiskerByMAQ1823AD39DT234AB532063E128AX": {
        "name": "Box e Whisker by MAQ",
        "desc": "Gráfico de caixa (boxplot) para análise estatística de distribuição de dados.",
        "value": "Permite visualizar mediana, quartis, outliers e dispersão de preços ou volumes em uma única visualização compacta."
    },
    "ChicletSlicer1448559807354": {
        "name": "Chiclet Slicer",
        "desc": "Segmentação visual no formato de botões coloridos (chiclets) com suporte a imagens.",
        "value": "Oferece uma experiência de filtragem visualmente aprimorada em comparação aos slicers padrão, facilitando a navegação entre categorias."
    },
    "deneb7E15AEF80B9E4D4F8E12924291ECE89A": {
        "name": "Deneb",
        "desc": "Visual de código aberto baseado nas gramáticas de visualização Vega e Vega-Lite.",
        "value": "Possibilita criar visualizações totalmente customizadas que não existem nativamente no Power BI, com controle preciso sobre layout e interatividade."
    },
    "EnhancedScatterChart1443994985041": {
        "name": "Enhanced Scatter Chart",
        "desc": "Gráfico de dispersão aprimorado com suporte a imagens, formas e rótulos enriquecidos.",
        "value": "Amplia a capacidade analítica de correlação entre variáveis com recursos visuais adicionais como imagens de logos e tamanhos de bolhas personalizados."
    },
    "scatterChartByAkvelon6CFB0DAB29E746BE850B6DFD1E2789FD": {
        "name": "Scatter Chart by Akvelon",
        "desc": "Gráfico de dispersão com funcionalidades estendidas de clustering e destaque de pontos.",
        "value": "Enriquece a análise comparativa entre múltiplas dimensões simultâneas."
    },
    "STANDALONEdeneb7E15AEF80B9E4D4F8E12924291ECE89A": {
        "name": "Deneb Standalone",
        "desc": "Versão autônoma do visual Deneb, com pacote independente de dependências.",
        "value": "Garante compatibilidade e portabilidade das visualizações Vega/Vega-Lite em diferentes ambientes Power BI."
    },
}


# ==============================================================================
# LLM INTEGRATION (opcional — fallback silencioso para heurísticas)
# ==============================================================================

def call_llm(prompt: str, llm_config: dict):
    """Chama LLM configurado e retorna a resposta em texto ou None se falhar."""
    if not llm_config or llm_config.get("provider") == "Desabilitado":
        return None
    try:
        import urllib.request as _req
        import json as _json
        import ssl

        provider = llm_config.get("provider", "")
        api_key  = llm_config.get("api_key", "").strip()
        model    = llm_config.get("model", "").strip()

        if provider == "Ollama (Local)":
            # Ollama expõe API compatível com OpenAI em localhost
            url     = "http://localhost:11434/v1/chat/completions"
            model   = model or "llama3.2:3b"
            headers = {"Content-Type": "application/json"}
            payload = _json.dumps({
                "model":    model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "stream":   False,
            }).encode()
            request = _req.Request(url, data=payload, headers=headers, method="POST")
            # Ollama roda local — sem SSL
            with _req.urlopen(request, timeout=120) as resp:
                result = _json.loads(resp.read())
                return result["choices"][0]["message"]["content"]

        elif provider == "Azure OpenAI":
            endpoint = llm_config.get("endpoint", "").rstrip("/")
            model    = model or "gpt-4o"
            url      = f"{endpoint}/openai/deployments/{model}/chat/completions?api-version=2024-02-01"
            headers  = {"Content-Type": "application/json", "api-key": api_key}

        elif provider == "GitHub Models":
            # GitHub Models — API compatível com OpenAI, usa token GitHub
            # Token pode ser informado manualmente ou lido das variáveis de ambiente
            token = (api_key
                     or os.environ.get("GH_TOKEN", "")
                     or os.environ.get("GITHUB_TOKEN", "")).strip()
            model   = model or "gpt-4o"
            url     = "https://models.inference.ai.azure.com/chat/completions"
            headers = {"Content-Type": "application/json",
                       "Authorization": f"Bearer {token}"}

        else:  # OpenAI direto
            model   = model or "gpt-4o"
            url     = "https://api.openai.com/v1/chat/completions"
            headers = {"Content-Type": "application/json",
                       "Authorization": f"Bearer {api_key}"}

        if provider not in ("Ollama (Local)",):
            payload = _json.dumps({
                "messages":   [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens":  800,
                "model":       model,
            }).encode()
            request = _req.Request(url, data=payload, headers=headers, method="POST")
            ctx = ssl.create_default_context()
            with _req.urlopen(request, context=ctx, timeout=60) as resp:
                result = _json.loads(resp.read())
                return result["choices"][0]["message"]["content"]

    except Exception:
        return None


def enhance_page_with_llm(page: dict, deep: dict, cls: dict, llm_config: dict):
    """
    Enriquece a narrativa de uma página usando LLM.
    Retorna dict com objetivo, resumo e leitura, ou None se LLM falhar/não configurado.
    """
    if not llm_config or llm_config.get("provider") == "Desabilitado":
        return None

    page_name = page["name"]
    profile   = cls.get("profile", "análise estratégica")

    kpis          = deep.get("kpis", [])
    filters       = deep.get("filters", [])
    chart_details = deep.get("chart_details", [])
    table_cols    = deep.get("table_cols", [])
    model_tables  = sorted({t for v in page["visuals"] for t in v.get("tables", []) if t})

    visual_composition = (
        f"{cls.get('card_count', 0)} cards, "
        f"{cls.get('chart_count', 0)} gráficos, "
        f"{cls.get('slicer_count', 0)} slicers, "
        f"{cls.get('content_count', 0)} total analíticos"
    )
    kpis_str   = ", ".join(kpis[:8]) if kpis else "—"
    charts_str = "; ".join(
        f"{c['type']}: {', '.join(c['measures'][:2])} × {', '.join(c['dims'][:2])}"
        for c in chart_details[:4]
    ) if chart_details else "—"
    tables_str    = " | ".join(", ".join(cols[:5]) for cols in table_cols[:2]) if table_cols else "—"
    filters_str   = ", ".join(filters[:6]) if filters else "—"
    model_tbl_str = ", ".join(model_tables[:8]) if model_tables else "—"

    prompt = (
        f'Você é um analista de BI sênior explicando um relatório Power BI para alguém que vai dar manutenção nele pela primeira vez.\n\n'
        f'Com base nos dados da aba abaixo, escreva em português corporativo e didático:\n\n'
        f'1. OBJETIVO: 1-2 frases sobre o propósito desta página — qual pergunta de negócio ela responde e quem a utiliza\n'
        f'2. RESUMO: 2-4 parágrafos explicando: o que esta aba analisa, quais indicadores são mais importantes, '
        f'como as informações se relacionam entre si, e o que uma variação nos números significaria para o negócio\n'
        f'3. LEITURA: 4-6 passos numerados ensinando como navegar a página — onde olhar primeiro, '
        f'como usar os filtros, o que cada grupo de visuais revela, e como interpretar os resultados\n\n'
        f'DADOS DA ABA "{page_name}":\n'
        f'- Visuais analíticos: {visual_composition}\n'
        f'- KPIs (cards): {kpis_str}\n'
        f'- Gráficos: {charts_str}\n'
        f'- Tabelas/Matrizes colunas: {tables_str}\n'
        f'- Filtros (slicers): {filters_str}\n'
        f'- Tabelas do modelo: {model_tbl_str}\n'
        f'- Perfil detectado: {profile}\n\n'
        f'Responda APENAS no formato JSON:\n'
        f'{{"objetivo": "...", "resumo": "...", "leitura": "..."}}'
    )

    raw = call_llm(prompt, llm_config)
    if not raw:
        return None

    try:
        text = raw.strip()
        if "```" in text:
            text = re.sub(r"```(?:json)?", "", text).strip().rstrip("`").strip()
        data = json.loads(text)
        return {
            "objetivo": str(data.get("objetivo", "")),
            "resumo":   str(data.get("resumo",   "")),
            "leitura":  str(data.get("leitura",  "")),
        }
    except Exception:
        return None


# ==============================================================================
# CAMADA 1: PARSING
# ==============================================================================

def extract_pbix(pbix_path: str) -> str:
    """Extrai o .pbix para um diretório temporário e retorna o caminho."""
    tmp = tempfile.mkdtemp(prefix="pbix_analyzer_")
    with zipfile.ZipFile(pbix_path, "r") as z:
        z.extractall(tmp)
    return tmp


def read_utf16le_json(path: str):
    """Lê um arquivo JSON com encoding UTF-16 LE (padrão Power BI)."""
    with open(path, "r", encoding="utf-16-le") as f:
        return json.load(f)


def read_utf8_json(path: str):
    """Lê um arquivo JSON com encoding UTF-8."""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def safe_json_loads(s):
    """Parseia JSON com tratamento seguro de erros."""
    if not s:
        return {}
    try:
        return json.loads(s)
    except Exception:
        return {}


def get_visual_label(vtype: str) -> str:
    return VISUAL_TYPE_LABELS.get(vtype, vtype)


def parse_select_fields(select_list):
    """Extrai campos/medidas de um prototypeQuery.Select."""
    fields = []
    for item in select_list:
        display = item.get("NativeReferenceName") or item.get("Name", "")
        if "Measure" in item:
            kind  = "Medida"
            table = item["Measure"]["Expression"].get("SourceRef", {}).get("Entity", "")
            col   = item["Measure"].get("Property", "")
        elif "Column" in item:
            kind  = "Coluna"
            table = item["Column"]["Expression"].get("SourceRef", {}).get("Entity", "")
            col   = item["Column"].get("Property", "")
        elif "Aggregation" in item:
            kind  = "Agregação"
            inner = item["Aggregation"].get("Expression", {})
            if "Column" in inner:
                table = inner["Column"]["Expression"].get("SourceRef", {}).get("Entity", "")
                col   = inner["Column"].get("Property", "")
            else:
                table, col = "", ""
        else:
            kind, table, col = "Campo", "", ""
        fields.append({"display": display, "kind": kind, "table": table, "column": col})
    return fields


def parse_visual(vc: dict) -> dict:
    """Parseia um visualContainer e retorna informações estruturadas."""
    config = safe_json_loads(vc.get("config", "{}"))
    sv     = config.get("singleVisual", {})
    vtype  = sv.get("visualType", "unknown")

    title = ""
    try:
        objects   = sv.get("objects", {})
        title_cfg = objects.get("title", [{}])[0].get("properties", {}).get("text", {})
        title     = title_cfg.get("expr", {}).get("Literal", {}).get("Value", "").strip("'")
    except Exception:
        pass

    pq          = sv.get("prototypeQuery", {})
    fields      = parse_select_fields(pq.get("Select", []))
    tables_used = list({f.get("Entity", "") for f in pq.get("From", []) if f.get("Entity")})

    filters_raw = safe_json_loads(vc.get("filters", "[]"))
    filter_cols = []
    for flt in (filters_raw if isinstance(filters_raw, list) else []):
        expr = flt.get("expression", {})
        col  = expr.get("Column", {}).get("Property", "")
        tbl  = expr.get("Column", {}).get("Expression", {}).get("SourceRef", {}).get("Entity", "")
        if col:
            filter_cols.append(f"{tbl}.{col}" if tbl else col)

    return {
        "type":    vtype,
        "label":   get_visual_label(vtype),
        "title":   title,
        "fields":  fields,
        "tables":  tables_used,
        "filters": filter_cols,
        "x":       vc.get("x", 0),
        "y":       vc.get("y", 0),
        "width":   vc.get("width", 0),
        "height":  vc.get("height", 0),
    }


def parse_layout(extract_dir: str) -> dict:
    """Parseia o arquivo Report/Layout e retorna a estrutura do relatório."""
    layout_path = os.path.join(extract_dir, "Report", "Layout")
    data = read_utf16le_json(layout_path)

    report    = {"theme": data.get("theme", ""), "pages": []}
    pages_raw = sorted(data.get("sections", []), key=lambda x: x.get("ordinal", 0))

    for page in pages_raw:
        filters_raw  = safe_json_loads(page.get("filters", "[]")) if isinstance(page.get("filters"), str) else page.get("filters", [])
        page_filters = []
        for flt in (filters_raw if isinstance(filters_raw, list) else []):
            expr = flt.get("expression", {})
            col  = expr.get("Column", {}).get("Property", "")
            tbl  = expr.get("Column", {}).get("Expression", {}).get("SourceRef", {}).get("Entity", "")
            if col:
                page_filters.append(f"{tbl}.{col}" if tbl else col)

        visuals    = [parse_visual(vc) for vc in page.get("visualContainers", [])]
        name_lower = page.get("displayName", "").lower()
        page_name  = page.get("name", "").lower()
        is_tooltip = "tooltip" in name_lower or "tooltip" in page_name

        report["pages"].append({
            "name":       page.get("displayName", "Sem nome"),
            "ordinal":    page.get("ordinal", 0),
            "width":      page.get("width", 1280),
            "height":     page.get("height", 720),
            "filters":    page_filters,
            "visuals":    visuals,
            "is_tooltip": is_tooltip,
        })

    return report


def parse_diagram_layout(extract_dir: str) -> list:
    """Retorna lista de nomes de tabelas do DiagramLayout."""
    path = os.path.join(extract_dir, "DiagramLayout")
    if not os.path.exists(path):
        return []
    try:
        data = read_utf16le_json(path)
    except Exception:
        return []
    tables = []
    for diagram in data.get("diagrams", []):
        for node in diagram.get("nodes", []):
            name = node.get("nodeIndex")
            if name and name not in tables:
                tables.append(name)
    return sorted(tables)


def infer_relationships(report: dict, tables: list) -> list:
    """
    Infere relacionamentos prováveis entre tabelas com base em:
    1. Co-ocorrência de tabelas no mesmo visual (alta confiança)
    2. Similaridade de nomes de colunas entre tabelas (média confiança)
    Retorna lista de dicts: {from_table, to_table, via, confidence, pages}
    """
    # Mapa: tabela → lista de campos com nomes
    table_fields: dict[str, set] = defaultdict(set)
    # Mapa: (tabelaA, tabelaB) → {pages, count}
    co_usage: dict[tuple, dict] = defaultdict(lambda: {"pages": set(), "count": 0, "fields": set()})

    for page in report["pages"]:
        for vis in page["visuals"]:
            tables_in_vis = [t for t in vis["tables"] if t]
            # Campos por tabela neste visual
            for fld in vis["fields"]:
                if fld.get("table") and fld.get("display"):
                    table_fields[fld["table"]].add(fld["display"].lower())
            # Co-ocorrência
            for i, ta in enumerate(tables_in_vis):
                for tb in tables_in_vis[i+1:]:
                    if ta == tb:
                        continue
                    key = tuple(sorted([ta, tb]))
                    co_usage[key]["pages"].add(page["name"])
                    co_usage[key]["count"] += 1
                    shared = {f["display"] for f in vis["fields"]
                              if f.get("table") == ta and f.get("display")}
                    co_usage[key]["fields"].update(shared)

    # Build inferred relationships from co-usage
    relationships = []
    seen = set()
    for (ta, tb), info in sorted(co_usage.items(), key=lambda x: -x[1]["count"]):
        if info["count"] < 1:
            continue
        key = (ta, tb)
        if key in seen:
            continue
        seen.add(key)

        # Heuristic: find shared or similar column names
        fields_a = table_fields.get(ta, set())
        fields_b = table_fields.get(tb, set())
        shared_cols = fields_a & fields_b
        # Also look for partial matches (e.g. "id" in both)
        id_cols_a = {f for f in fields_a if any(k in f for k in ["id","key","cod","code","_sk","pk"])}
        id_cols_b = {f for f in fields_b if any(k in f for k in ["id","key","cod","code","_sk","pk"])}
        likely_join = shared_cols or (id_cols_a & id_cols_b)

        via_str = ", ".join(sorted(likely_join)[:2]) if likely_join else "co-uso de campos"
        confidence = "Alta" if info["count"] >= 3 and likely_join else (
                     "Média" if info["count"] >= 2 or likely_join else "Baixa")

        relationships.append({
            "from_table":  ta,
            "to_table":    tb,
            "via":         via_str,
            "confidence":  confidence,
            "count":       info["count"],
            "pages":       sorted(info["pages"])[:3],
        })

    # Sort: high confidence first, then by count
    conf_order = {"Alta": 0, "Média": 1, "Baixa": 2}
    relationships.sort(key=lambda r: (conf_order[r["confidence"]], -r["count"]))
    return relationships


def generate_relationship_diagram(relationships: list, tables: list,
                                   table_classifications: dict) -> bytes | None:
    """
    Gera um diagrama de rede de relacionamentos como imagem PNG em memória.
    Retorna bytes da imagem ou None se matplotlib/networkx não disponível.
    """
    try:
        import io
        import networkx as nx
        import matplotlib
        matplotlib.use("Agg")          # backend sem janela
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
    except ImportError:
        return None

    if not relationships:
        return None

    G = nx.Graph()

    # Paleta de tipos de tabela
    TYPE_COLORS = {
        "fato":        "#FF8C00",   # laranja
        "dimensão":    "#FFBE00",   # amarelo Construction
        "calendário":  "#4DAAFF",   # azul
        "medidas":     "#A0A0A0",   # cinza
        "geográfica":  "#66CC66",   # verde suave
        "mapeamento":  "#CC99FF",   # lilás
        "auxiliar":    "#DDDDDD",   # cinza claro
    }
    DEFAULT_COLOR = "#EEEEEE"

    def table_color(tname):
        tc = table_classifications.get(tname, {}).get("type", "auxiliar")
        for k, c in TYPE_COLORS.items():
            if k in tc.lower():
                return c
        return DEFAULT_COLOR

    # Adiciona nós
    all_tbl = set()
    for r in relationships:
        all_tbl.add(r["from_table"])
        all_tbl.add(r["to_table"])
    for t in all_tbl:
        G.add_node(t)

    # Adiciona arestas com peso
    conf_w = {"Alta": 3.0, "Média": 1.8, "Baixa": 0.8}
    for r in relationships:
        G.add_edge(r["from_table"], r["to_table"],
                   weight=conf_w[r["confidence"]],
                   confidence=r["confidence"],
                   via=r["via"])

    # Layout
    pos = nx.spring_layout(G, seed=42, k=2.5 / max(len(G.nodes())**0.5, 1))

    fig, ax = plt.subplots(figsize=(14, 9), facecolor="#1A1A1A")
    ax.set_facecolor("#1A1A1A")

    # Desenhar arestas por confiança
    for conf, style, alpha in [("Alta", "solid", 0.9),
                                 ("Média", "dashed", 0.65),
                                 ("Baixa", "dotted", 0.4)]:
        edges = [(u, v) for u, v, d in G.edges(data=True) if d["confidence"] == conf]
        widths = [conf_w[conf] * 1.2 for _ in edges]
        nx.draw_networkx_edges(G, pos, edgelist=edges, width=widths,
                               edge_color="#FFBE00", style=style, alpha=alpha, ax=ax)

    # Desenhar nós
    node_colors = [table_color(n) for n in G.nodes()]
    node_sizes  = [max(1400, 250 * len(n)) for n in G.nodes()]
    nx.draw_networkx_nodes(G, pos, node_color=node_colors, node_size=node_sizes,
                           linewidths=1.5, edgecolors="#FFBE00", ax=ax)

    # Labels dos nós
    nx.draw_networkx_labels(G, pos, font_size=7.5, font_color="#1A1A1A",
                            font_weight="bold", ax=ax)

    # Legenda de tipos
    legend_items = [
        mpatches.Patch(color=c, label=t.capitalize())
        for t, c in TYPE_COLORS.items()
        if any(t in (table_classifications.get(n, {}).get("type", "") or "").lower()
               for n in G.nodes())
    ]
    legend_items += [
        mpatches.Patch(color="none", label=""),
        plt.Line2D([0],[0], color="#FFBE00", linewidth=2.5, linestyle="solid",  label="Confiança Alta"),
        plt.Line2D([0],[0], color="#FFBE00", linewidth=1.8, linestyle="dashed", label="Confiança Média"),
        plt.Line2D([0],[0], color="#FFBE00", linewidth=0.8, linestyle="dotted", label="Confiança Baixa"),
    ]
    if legend_items:
        ax.legend(handles=legend_items, loc="upper left", fontsize=7,
                  facecolor="#2A2A2A", edgecolor="#FFBE00",
                  labelcolor="white", framealpha=0.85)

    ax.set_title("Mapa de Relacionamentos entre Tabelas\n(inferido a partir do uso em visuais)",
                 color="#FFBE00", fontsize=11, pad=12, fontweight="bold")
    ax.axis("off")
    plt.tight_layout(pad=1.0)

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="#1A1A1A")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def parse_connections(extract_dir: str) -> dict:
    """Lê o arquivo Connections do .pbix."""
    path = os.path.join(extract_dir, "Connections")
    if not os.path.exists(path):
        return {}
    try:
        return read_utf8_json(path)
    except Exception:
        return {}


def parse_metadata(extract_dir: str) -> dict:
    """Lê o arquivo Metadata do .pbix."""
    path = os.path.join(extract_dir, "Metadata")
    if not os.path.exists(path):
        return {}
    try:
        return read_utf8_json(path)
    except Exception:
        return {}


# ==============================================================================
# CAMADA 2: INTELIGÊNCIA / ANÁLISE
# ==============================================================================

def classify_field(name: str, kind: str) -> str:
    """Classifica um campo em uma categoria analítica com base em heurísticas."""
    n = name.lower()
    if kind == "Medida":
        return "KPI/Medida"
    temporal_kw = ["data", "mes", "ano", "semana", "trimestre", "quarter", "year",
                   "month", "week", "dia", "date", "periodo", "período"]
    geo_kw = ["pais", "estado", "regiao", "cidade", "uf", "geo", "localidade",
              "municipio", "país", "região"]
    prod_kw = ["produto", "modelo", "familia", "categoria", "versao", "serie",
               "sku", "item", "família", "versão", "série"]
    com_kw  = ["dealer", "concession", "canal", "cliente", "empresa", "fornecedor",
               "parceiro", "revend"]
    if any(w in n for w in temporal_kw):
        return "Dimensão temporal"
    if any(w in n for w in geo_kw):
        return "Dimensão geográfica"
    if any(w in n for w in prod_kw):
        return "Dimensão de produto"
    if any(w in n for w in com_kw):
        return "Dimensão comercial"
    if any(w in n for w in ["slicer", "filtro", "segmenta", "flag", "tipo", "classe",
                             "grupo", "status"]):
        return "Filtro/Segmentação"
    if kind == "Coluna":
        return "Dimensão comercial"
    return "Campo auxiliar"


def classify_table(name: str, pages_used: list) -> dict:
    """Classifica uma tabela do modelo com base no nome e padrão de uso."""
    n = name.lower()
    if any(w in n for w in ["calendar", "calendario", "calendário", "date", "datas",
                             "tempo", "dt_"]):
        return {
            "type": "calendário",
            "role": "Tabela de dimensão temporal usada para filtros por data, mês, trimestre e ano."
        }
    if any(w in n for w in ["medidas", "measure", "kpi", "_medidas", "calculations",
                             "métricas", "metricas"]):
        return {
            "type": "medidas",
            "role": "Tabela de medidas DAX (sem dados brutos), concentra os indicadores calculados do relatório."
        }
    if any(w in n for w in ["depara", "de_para", "de para", "mapeamento", "mapping",
                             "crosswalk", "lookup"]):
        return {
            "type": "mapeamento/de-para",
            "role": "Tabela de mapeamento entre códigos ou nomes diferentes para a mesma entidade."
        }
    if any(w in n for w in ["pais", "estado", "regiao", "geo", "localidade",
                             "municipio", "países", "país"]):
        return {
            "type": "geográfica",
            "role": "Tabela de dimensão geográfica com estados, regiões ou países."
        }
    if any(w in n for w in ["pesquisa", "fato", "fact", "vendas", "venda", "preco",
                             "price", "transacao", "transação", "precificacao"]):
        return {
            "type": "fato",
            "role": "Tabela de fatos com dados transacionais ou de pesquisa, base primária da análise."
        }
    if n.startswith("db_") or any(w in n for w in ["dim", "dimension"]):
        return {
            "type": "dimensão",
            "role": "Tabela de dimensão que suporta filtros e agrupamentos nas análises."
        }
    if any(w in n for w in ["status", "reason", "tipo", "classe", "grupo",
                             "parametro", "parâmetro", "config"]):
        return {
            "type": "auxiliar",
            "role": "Tabela auxiliar de suporte com categorias ou configurações para segmentação."
        }
    if len(pages_used) >= 2:
        return {
            "type": "dimensão",
            "role": "Tabela referenciada em múltiplas páginas, provável dimensão compartilhada."
        }
    return {
        "type": "auxiliar",
        "role": "Tabela auxiliar ou de suporte com papel complementar no modelo."
    }


def classify_page(page: dict) -> dict:
    """Classifica uma página com base no conteúdo real — sem checks hardcoded de nome."""
    name    = page["name"].lower()
    visuals = page["visuals"]
    total   = len(visuals)

    content_v    = [v for v in visuals if v["type"] not in DECORATION_TYPES]
    slicers      = [v for v in visuals if "slicer" in v["type"].lower() or v["type"] in SLICER_TYPES]
    cards        = [v for v in visuals if v["type"] in ("card", "multiRowCard")]
    matrices     = [v for v in visuals if v["type"] in ("pivotTable", "tableEx")]
    maps         = [v for v in visuals if "map" in v["type"].lower()]
    charts       = [v for v in visuals if any(w in v["type"].lower() for w in
                    ["chart", "line", "column", "bar", "donut", "pie"])]
    stats_v      = [v for v in visuals if ("box" in v["type"].lower() or
                    "BoxandWhisker" in v["type"] or
                    "scatter" in v["type"].lower() or "Scatter" in v["type"])]

    has_maps         = len(maps) > 0
    has_stats        = len(stats_v) > 0
    has_matrices     = len(matrices) > 0
    has_many_slicers = len(slicers) > 4

    # Campos e tabelas desta página para detecção de tema
    all_field_names  = " ".join(f["display"].lower() for v in visuals for f in v["fields"] if f["display"])
    all_table_names  = " ".join(t.lower() for v in visuals for t in v.get("tables", []))
    combined         = all_field_names + " " + all_table_names + " " + name

    # ── Sinais de tema (sem termos ultra-específicos como "apn", "tabela") ────
    sig_price   = any(w in combined for w in ["preco", "preço", "price", "pricing", "lista"])
    sig_disc    = any(w in combined for w in ["desconto", "discount", "desvio", "variação", "diferença"])
    sig_dealer  = any(w in combined for w in ["dealer", "concession", "revendedor", "canal de venda", "tsm"])
    sig_volume  = any(w in combined for w in ["volume", "qtd", "quantidade", "unidade", "venda"])
    sig_compet  = any(w in combined for w in ["competi", "concorren", "mercado", "rival", "concorrente"])
    sig_geo     = any(w in combined for w in ["estado", "regiao", "região", "uf", "pais", "país", "cidade"])
    sig_product = any(w in combined for w in ["modelo", "produto", "familia", "família", "sku", "categoria"])
    sig_aprov   = any(w in combined for w in ["aprova", "status", "aprovação", "pendente", "rejeit", "formulari"])
    sig_fin     = any(w in combined for w in ["receita", "custo", "margem", "lucro", "budget", "meta", "orçamento"])

    # ── Tema dominante derivado dos sinais ────────────────────────────────────
    if sig_price and sig_compet:
        dominant_theme = "comparação competitiva de preços"
    elif sig_price and sig_disc:
        dominant_theme = "análise de preços e descontos"
    elif sig_price and sig_dealer:
        dominant_theme = "preços e desempenho de canais"
    elif sig_price:
        dominant_theme = "análise de preços"
    elif sig_fin:
        dominant_theme = "análise financeira e de metas"
    elif sig_aprov:
        dominant_theme = "acompanhamento de aprovações e status"
    elif sig_dealer and sig_volume:
        dominant_theme = "desempenho comercial de dealers e volumes"
    elif sig_dealer:
        dominant_theme = "desempenho de dealers e canais"
    elif sig_volume and sig_compet:
        dominant_theme = "volume e posicionamento competitivo"
    elif sig_volume:
        dominant_theme = "análise de volumes e quantidades"
    elif sig_product:
        dominant_theme = "análise de portfólio de produtos"
    elif sig_geo:
        dominant_theme = "análise geográfica e territorial"
    else:
        # Fallback: usa as palavras significativas do nome da página
        words = [w for w in page["name"].split() if len(w) > 3]
        dominant_theme = f"análise de {' '.join(words[:3]).lower()}" if words else "análise de dados"

    # ── Perfil: keywords genéricas + composição de visuais ───────────────────
    home_kw = ["home", "início", "inicio", "menu", "navegação", "navegacao",
               "apresentação", "capa", "portal", "landing"]
    ger_kw  = ["gerencial", "gestão", "gestao", "executivo", "diretoria",
               "dashboard", "overview", "resumo", "sumário", "consolidado"]
    op_kw   = ["operacional", "operação", "operacao", "detalhe", "detalhad",
               "acompanhamento", "monitoramento", "tracking", "lista"]

    if any(w in name for w in home_kw):
        profile = "home/navegação"
    elif any(w in name for w in ger_kw):
        profile = "análise gerencial"
    elif any(w in name for w in op_kw):
        profile = "análise operacional"
    elif has_stats:
        profile = "análise exploratória"
    elif len(cards) >= 4 and len(charts) <= 2:
        profile = "análise gerencial"
    elif has_many_slicers or has_matrices:
        profile = "análise operacional"
    else:
        profile = "análise estratégica"

    # ── Campos reais para usar nas narrativas ─────────────────────────────────
    slicer_fld_names = list(dict.fromkeys(
        f["display"] for v in slicers for f in v["fields"] if f["display"]
    ))
    card_fld_names = list(dict.fromkeys(
        f["display"] for v in cards for f in v["fields"] if f["display"]
    ))
    chart_type_labels = list(dict.fromkeys(get_visual_label(v["type"]) for v in charts))

    # ── Key insights: baseados no conteúdo real ────────────────────────────────
    key_insights = []
    if len(cards) >= 2:
        kpi_str = f": {', '.join(card_fld_names[:3])}" if card_fld_names else ""
        key_insights.append(f"Apresenta {len(cards)} KPIs em cartão para leitura rápida{kpi_str}.")
    if len(slicers) >= 2:
        slicer_str = f" por {', '.join(slicer_fld_names[:3])}" if slicer_fld_names else ""
        key_insights.append(f"Conta com {len(slicers)} filtros interativos — recortes{slicer_str}.")
    if has_maps:
        key_insights.append("Inclui visualização geográfica para análise territorial.")
    if has_stats:
        key_insights.append("Contém análise estatística (Box & Whisker / Dispersão) para distribuição e outliers.")
    if has_matrices:
        key_insights.append(f"Disponibiliza {len(matrices)} tabela(s)/matriz(es) para drill-down detalhado.")
    if len(charts) >= 3:
        chart_str = f" ({', '.join(chart_type_labels[:2])})" if chart_type_labels else ""
        key_insights.append(f"Utiliza {len(charts)} gráficos analíticos{chart_str} para comparação e tendências.")
    if not key_insights:
        key_insights.append(f"{len(content_v)} elementos analíticos para visualização de {dominant_theme}.")

    # ── Objetivo: construído a partir dos elementos reais da página ───────────
    page_display = page["name"]
    vis_parts = []
    if cards:
        vis_parts.append(f"{len(cards)} indicador(es)-chave")
    if charts:
        vis_parts.append(f"{len(charts)} gráfico(s) analítico(s)")
    if matrices:
        vis_parts.append(f"{len(matrices)} tabela(s)/matriz(es)")
    if maps:
        vis_parts.append("visualização geográfica")
    if stats_v:
        vis_parts.append("análise estatística")
    vis_str = ", ".join(vis_parts) if vis_parts else f"{len(content_v)} visuais analíticos"

    filter_str = ""
    if slicer_fld_names:
        filter_str = f" Os filtros disponíveis permitem recortes por {', '.join(slicer_fld_names[:4])}."

    if profile == "home/navegação":
        objective = (
            f"Hub de navegação e entrada do relatório '{page_display}'. "
            f"Apresenta os principais módulos analíticos e contextualiza o escopo do relatório."
        )
    else:
        objective = f"A página '{page_display}' reúne {vis_str} para {dominant_theme}.{filter_str}"

    # ── Guia de leitura: referencia campos e visuais reais ────────────────────
    guide_parts = []
    if profile == "home/navegação":
        guide_parts.append(
            "Utilize os botões e links de navegação para acessar os módulos analíticos. "
            "Os indicadores de resumo contextualizam período e escopo da análise."
        )
    else:
        if slicer_fld_names:
            guide_parts.append(
                f"Inicie aplicando os filtros ({', '.join(slicer_fld_names[:4])}) "
                f"para delimitar período, segmento ou contexto de interesse."
            )
        if card_fld_names:
            guide_parts.append(
                f"Leia os cartões de KPI ({', '.join(card_fld_names[:4])}) "
                f"para obter a visão consolidada dos indicadores principais."
            )
        elif cards:
            guide_parts.append(
                f"Os {len(cards)} cartões de KPI oferecem a visão consolidada dos indicadores desta perspectiva."
            )
        if chart_type_labels:
            guide_parts.append(
                f"Os gráficos ({', '.join(chart_type_labels[:2])}) permitem análise comparativa "
                f"e identificação de tendências em {dominant_theme}."
            )
        if matrices:
            guide_parts.append(
                "As tabelas e matrizes permitem aprofundamento por dimensão — "
                "ideal para investigar valores específicos e anomalias."
            )
        if has_stats:
            guide_parts.append(
                "Os gráficos estatísticos revelam distribuição, variabilidade e outliers, "
                "úteis para análise de concentração e desvios significativos."
            )
        if has_maps:
            guide_parts.append("O mapa permite identificar padrões geográficos e concentrações por território.")
        if not guide_parts:
            guide_parts.append(
                f"Explore os {len(content_v)} visuais disponíveis para análise de {dominant_theme}, "
                f"utilizando os filtros para segmentar conforme necessidade."
            )

    reading_guide = " ".join(guide_parts)

    return {
        "profile":           profile,
        "objective":         objective,
        "key_insights":      key_insights[:4],
        "reading_guide":     reading_guide,
        "dominant_theme":    dominant_theme,
        "has_maps":          has_maps,
        "has_stats_visuals": has_stats,
        "has_matrices":      has_matrices,
        "has_many_slicers":  has_many_slicers,
        "slicer_count":      len(slicers),
        "card_count":        len(cards),
        "chart_count":       len(charts),
        "content_count":     len(content_v),
        "decoration_count":  total - len(content_v),
        "slicer_fields":     slicer_fld_names,
        "card_fields":       card_fld_names,
    }


def analyze_page_deeply(page: dict) -> dict:
    """
    Lê cada visual da página individualmente e extrai contexto rico:
    - KPIs dos cards (nomes reais)
    - Filtros dos slicers (campos reais)
    - O que cada gráfico analisa (medida × dimensão)
    - Colunas detalhadas nas tabelas/matrizes
    - Análises estatísticas (quais variáveis)
    - Análises geográficas (medida × dimensão geo)
    Tudo 100% dinâmico — sem nada fixo no código.
    """
    kpis          = []   # nomes dos indicadores em cards
    filters       = []   # campos dos slicers
    chart_details = []   # {"type": label, "measures": [...], "dims": [...]}
    table_cols    = []   # lista de listas de colunas por tabela/matriz
    map_details   = []   # {"measure": str, "geo": str}
    stat_details  = []   # {"type": label, "vars": [...]}

    for v in page["visuals"]:
        vtype  = v["type"]
        fields = v["fields"]

        if vtype in DECORATION_TYPES:
            continue

        # Separa medidas/agregações de dimensões/colunas
        measures   = [f["display"] for f in fields
                      if f["kind"] in ("Medida", "Agregação") and f["display"]]
        dimensions = [f["display"] for f in fields
                      if f["kind"] == "Coluna" and f["display"]]
        all_fnames = [f["display"] for f in fields if f["display"]]

        # ── Cards ──────────────────────────────────────────────────────────
        if vtype in ("card", "multiRowCard"):
            for nm in all_fnames:
                if nm not in kpis:
                    kpis.append(nm)

        # ── Slicers ────────────────────────────────────────────────────────
        elif vtype in SLICER_TYPES or "slicer" in vtype.lower():
            for nm in all_fnames:
                if nm not in filters:
                    filters.append(nm)

        # ── Tabelas / Matrizes ─────────────────────────────────────────────
        elif vtype in ("pivotTable", "tableEx"):
            cols = list(dict.fromkeys(all_fnames))
            if cols:
                table_cols.append(cols)

        # ── Mapas ──────────────────────────────────────────────────────────
        elif "map" in vtype.lower() or "Map" in vtype:
            geo  = next((d for d in dimensions if d), None)
            meas = next((m for m in measures  if m), None)
            if geo or meas:
                map_details.append({
                    "measure": meas or "—",
                    "geo":     geo  or "—",
                })

        # ── Estatísticos (Box & Whisker, Scatter) ──────────────────────────
        elif ("box" in vtype.lower() or "BoxandWhisker" in vtype
              or "scatter" in vtype.lower() or "Scatter" in vtype):
            vars_ = list(dict.fromkeys(all_fnames))
            if vars_:
                stat_details.append({
                    "type": get_visual_label(vtype),
                    "vars": vars_,
                })

        # ── Gráficos: qualquer visual com medidas/dimensões que não seja
        #    card, slicer, tabela, mapa ou estatístico (catch-all)
        elif (measures or dimensions) and vtype not in DECORATION_TYPES:
            chart_details.append({
                "type":     get_visual_label(vtype),
                "measures": measures[:3],
                "dims":     dimensions[:3],
            })

    return {
        "kpis":          list(dict.fromkeys(kpis)),
        "filters":       list(dict.fromkeys(filters)),
        "chart_details": chart_details,
        "table_cols":    table_cols,
        "map_details":   map_details,
        "stat_details":  stat_details,
    }


def generate_page_narrative(page: dict, cls: dict) -> dict:
    """
    Gera narrativa executiva completa por página, lendo o conteúdo real de
    cada visual — 100% dinâmico, sem nenhum nome de página ou campo fixo.
    """
    name    = page["name"]
    profile = cls.get("profile", "análise estratégica")
    theme   = cls.get("dominant_theme", "dados")

    # Leitura profunda: cada visual em detalhe
    deep = analyze_page_deeply(page)
    kpis          = deep["kpis"]
    filters       = deep["filters"]
    chart_details = deep["chart_details"]
    table_cols    = deep["table_cols"]
    map_details   = deep["map_details"]
    stat_details  = deep["stat_details"]

    # Contagens rápidas
    n_content = cls.get("content_count", 0)

    # ── Resumo executivo ──────────────────────────────────────────────────────
    if profile == "home/navegação":
        exec_sum = (
            f"A página '{name}' funciona como hub de navegação e ponto de entrada do relatório. "
            f"Apresenta os módulos analíticos disponíveis e contextualiza escopo e período da análise."
        )
    else:
        sentences = []

        # Frase de abertura com KPIs reais
        if kpis:
            kpi_str = ", ".join(kpis[:5])
            extra   = f" e mais {len(kpis) - 5}" if len(kpis) > 5 else ""
            sentences.append(
                f"A página '{name}' expõe os indicadores {kpi_str}{extra}, "
                f"consolidados como referência central para {theme}."
            )
        else:
            sentences.append(
                f"A página '{name}' apresenta {n_content} visuais analíticos "
                f"voltados para {theme}."
            )

        # Gráficos: descreve o que cada um analisa
        chart_sentences = []
        for ch in chart_details[:4]:
            if ch["measures"] and ch["dims"]:
                chart_sentences.append(
                    f"{ch['type']} cruzando {', '.join(ch['measures'][:2])} "
                    f"por {', '.join(ch['dims'][:2])}"
                )
            elif ch["measures"]:
                chart_sentences.append(
                    f"{ch['type']} exibindo {', '.join(ch['measures'][:2])}"
                )
            elif ch["dims"]:
                chart_sentences.append(
                    f"{ch['type']} por dimensão {', '.join(ch['dims'][:2])}"
                )
        if chart_sentences:
            sentences.append(
                f"A análise é suportada por: {'; '.join(chart_sentences)}."
            )

        # Tabelas/matrizes: lista as colunas reais
        for cols in table_cols[:2]:
            col_str = ", ".join(cols[:6])
            extra   = f" e mais {len(cols)-6} colunas" if len(cols) > 6 else ""
            sentences.append(
                f"A tabela detalha o dado ao nível de registro pelas colunas: "
                f"{col_str}{extra}."
            )

        # Análise estatística: diz quais variáveis
        for st in stat_details[:2]:
            var_str = ", ".join(st["vars"][:4])
            sentences.append(
                f"{st['type']} analisa a distribuição e variabilidade de: {var_str}."
            )

        # Mapa: diz o que está sendo mapeado
        for mp in map_details[:1]:
            sentences.append(
                f"Visualização geográfica mapeia {mp['measure']} por {mp['geo']}."
            )

        # Filtros: enumera os campos reais
        if filters:
            filt_str = ", ".join(filters[:6])
            extra    = f" (+{len(filters)-6} mais)" if len(filters) > 6 else ""
            sentences.append(
                f"A segmentação interativa cobre: {filt_str}{extra}."
            )

        exec_sum = " ".join(sentences)

    # ── Guia de leitura ───────────────────────────────────────────────────────
    if profile == "home/navegação":
        reading_guide = (
            "Utilize os botões e links de navegação para acessar os módulos analíticos. "
            "Os indicadores de resumo contextualizam período e escopo da análise."
        )
    else:
        steps = []

        # Passo 1 — aplicar filtros (com nomes reais)
        if filters:
            steps.append(
                f"1. Aplique os filtros ({', '.join(filters[:4])}) para delimitar "
                f"período, segmento ou contexto de interesse."
            )

        # Passo 2 — ler KPIs (com nomes reais)
        if kpis:
            steps.append(
                f"2. Leia os indicadores ({', '.join(kpis[:4])}) para obter "
                f"a visão consolidada dos valores-chave."
            )
        elif n_content > 0:
            steps.append("2. Consulte os cartões de KPI para a visão consolidada dos indicadores.")

        # Passo 3 — explorar gráficos (com o que analisam)
        if chart_details:
            first = chart_details[0]
            if first["measures"] and first["dims"]:
                steps.append(
                    f"3. Explore os gráficos para analisar como "
                    f"{', '.join(first['measures'][:2])} varia por "
                    f"{', '.join(first['dims'][:2])} e identifique tendências."
                )
            else:
                steps.append("3. Explore os gráficos para identificar tendências e comparações.")

        # Passo 4 — tabelas para drill-down
        if table_cols:
            cols_str = ", ".join(table_cols[0][:4])
            steps.append(
                f"4. Utilize a(s) tabela(s) com colunas ({cols_str}) "
                f"para aprofundamento ao nível de detalhe desejado."
            )

        # Passo 5 — estatística
        if stat_details:
            sv = stat_details[0]
            steps.append(
                f"5. Os gráficos estatísticos ({sv['type']}) sobre "
                f"{', '.join(sv['vars'][:3])} revelam distribuição e outliers."
            )

        # Passo 6 — mapa
        if map_details:
            mp = map_details[0]
            steps.append(
                f"6. O mapa mostra {mp['measure']} por {mp['geo']} — "
                f"ideal para identificar concentrações regionais."
            )

        if not steps:
            steps.append(
                f"Explore os {n_content} visuais disponíveis. "
                f"Utilize os filtros para segmentar a análise de {theme} conforme necessário."
            )

        reading_guide = " ".join(steps)

    return {
        "objective":         cls.get("objective", ""),
        "executive_summary": exec_sum.strip(),
        "reading_guide":     reading_guide,
        "deep":              deep,   # passa para o renderer usar na seção de KPIs/filtros
    }


def infer_business_entities(all_fields: dict, all_tables: list) -> list:
    """Infere as entidades de negócio presentes no relatório."""
    combined = " ".join(list(all_fields.keys()) + list(all_tables)).lower()
    entity_map = {
        "preço":          any(w in combined for w in ["preco", "preço", "price", "apn"]),
        "dealer/canal":   any(w in combined for w in ["dealer", "canal", "concession", "revendedor"]),
        "modelo/produto": any(w in combined for w in ["modelo", "produto", "familia", "sku", "versão"]),
        "competição":     any(w in combined for w in ["competi", "concorren", "mercado", "rival"]),
        "região/estado":  any(w in combined for w in ["estado", "regiao", "uf", "pais", "país"]),
        "volume":         any(w in combined for w in ["volume", "qtd", "quantidade", "unidade"]),
        "desconto":       any(w in combined for w in ["desconto", "discount", "desvio", "diferença"]),
        "tempo/período":  any(w in combined for w in ["data", "mes", "ano", "calendar", "periodo"]),
        "empresa":        any(w in combined for w in ["empresa", "empresometro", "cnpj"]),
    }
    return [entity for entity, found in entity_map.items() if found]


def group_fields_by_role(page_fields: list) -> dict:
    """Agrupa campos de uma página por papel analítico."""
    groups = {
        "KPIs":       [],
        "Dimensões":  [],
        "Filtros":    [],
        "Temporal":   [],
        "Geográfico": [],
        "Outros":     [],
    }
    for fld in page_fields:
        role = classify_field(fld["display"], fld["kind"])
        if role == "KPI/Medida":
            groups["KPIs"].append(fld)
        elif role in ("Dimensão comercial", "Dimensão de produto"):
            groups["Dimensões"].append(fld)
        elif role == "Filtro/Segmentação":
            groups["Filtros"].append(fld)
        elif role == "Dimensão temporal":
            groups["Temporal"].append(fld)
        elif role == "Dimensão geográfica":
            groups["Geográfico"].append(fld)
        else:
            groups["Outros"].append(fld)
    return {k: v for k, v in groups.items() if v}


def build_analysis(report: dict, tables: list, connections: dict, metadata: dict) -> dict:
    """Constrói a análise agregada com enriquecimentos heurísticos."""
    all_fields        = defaultdict(lambda: {"count": 0, "pages": set(), "kind": "Campo", "table": ""})
    all_tables_used   = defaultdict(lambda: {"pages": set()})
    visual_type_count = defaultdict(int)

    main_pages    = [p for p in report["pages"] if not p["is_tooltip"]]
    tooltip_pages = [p for p in report["pages"] if p["is_tooltip"]]

    for page in report["pages"]:
        for vis in page["visuals"]:
            visual_type_count[vis["type"]] += 1
            for tbl in vis["tables"]:
                if tbl:
                    all_tables_used[tbl]["pages"].add(page["name"])
            for fld in vis["fields"]:
                key = fld["display"]
                if key:
                    all_fields[key]["count"] += 1
                    all_fields[key]["pages"].add(page["name"])
                    all_fields[key]["kind"] = fld["kind"]
                    if fld.get("table"):
                        all_fields[key]["table"] = fld["table"]

    # Page classifications
    page_classifications = {}
    for page in report["pages"]:
        page_classifications[page["name"]] = classify_page(page)

    # Table set and classifications
    all_tables_set = set(tables)
    for tbl_name in all_tables_used:
        if tbl_name:
            all_tables_set.add(tbl_name)

    table_classifications = {}
    for tname in sorted(all_tables_set):
        pages = list(all_tables_used.get(tname, {}).get("pages", set()))
        table_classifications[tname] = classify_table(tname, pages)

    biz_entities = infer_business_entities(dict(all_fields), list(all_tables_set))

    used_custom   = {k: v for k, v in CUSTOM_VISUAL_MAP.items() if visual_type_count.get(k, 0) > 0}
    unused_custom = {k: v for k, v in CUSTOM_VISUAL_MAP.items() if visual_type_count.get(k, 0) == 0}

    # Inferir relacionamentos entre tabelas
    relationships = infer_relationships(report, list(all_tables_set))

    return {
        "main_pages":            main_pages,
        "tooltip_pages":         tooltip_pages,
        "all_fields":            dict(all_fields),
        "important_fields":      {k: v for k, v in all_fields.items()
                                  if v["count"] >= 2 or v["kind"] == "Medida"},
        "all_tables_used":       dict(all_tables_used),
        "all_tables_set":        sorted(all_tables_set),
        "visual_type_count":     dict(visual_type_count),
        "diagram_tables":        tables,
        "connections":           connections,
        "metadata":              metadata,
        "total_pages":           len(report["pages"]),
        "total_visuals":         sum(len(p["visuals"]) for p in report["pages"]),
        "page_classifications":  page_classifications,
        "table_classifications": table_classifications,
        "biz_entities":          biz_entities,
        "used_custom":           used_custom,
        "unused_custom":         unused_custom,
        "relationships":         relationships,
    }


# ==============================================================================
# CAMADA 3: RENDERIZAÇÃO DO DOCUMENTO WORD
# ==============================================================================

# ── Helpers de formatação XML/Word ────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Define cor de fundo de uma célula de tabela Word."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, color_hex="FFBE00", width_pt=12, sides=None):
    """Define bordas coloridas em uma célula."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    if sides is None:
        sides = ["top", "left", "bottom", "right"]
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(width_pt))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_paragraph_border_bottom(para, color_hex="FFBE00", size_pt=6):
    """Adiciona borda inferior a um parágrafo."""
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt * 8))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_header_row(table, texts: list, bg_hex="FFBE00", text_color=None):
    """Adiciona linha de cabeçalho formatada."""
    if text_color is None:
        text_color = JD_DARK   # preto sobre fundo amarelo Construction
    row = table.rows[0]
    for i, text in enumerate(texts):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        cell.text = ""
        set_cell_bg(cell, bg_hex)
        para           = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run            = para.add_run(str(text))
        run.font.bold  = True
        run.font.color.rgb = text_color
        run.font.size  = Pt(9)


def add_data_row(table, row_idx: int, values: list, alternate=False):
    """Adiciona linha de dados com zebra striping."""
    row = table.rows[row_idx]
    bg  = ALT_ROW_BG if alternate else "FFFFFF"
    for i, val in enumerate(values):
        if i >= len(row.cells):
            break
        cell       = row.cells[i]
        cell.text  = ""
        set_cell_bg(cell, bg)
        para = cell.paragraphs[0]
        run  = para.add_run(str(val) if val is not None else "—")
        run.font.size = Pt(9)


# ── Helpers de estrutura de documento ────────────────────────────────────────

def add_section_header(doc, number: int, title: str):
    """Adiciona cabeçalho de seção com fundo preto e texto amarelo."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, DARK_GREEN_BG)
    set_cell_border(cell, color_hex="FFBE00", width_pt=4)
    para           = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(f"  {number}.  {title.upper()}")
    run.font.bold      = True
    run.font.color.rgb = JD_YELLOW   # amarelo sobre preto
    run.font.size      = Pt(14)
    doc.add_paragraph()
    return tbl


def add_subsection_header(doc, title: str, level=2):
    """Adiciona sub-cabeçalho com cor e separador inferior."""
    p   = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold      = True
    run.font.size      = Pt(12 if level == 2 else 10)
    run.font.color.rgb = JD_DARK if level == 2 else DARK_GRAY
    set_paragraph_border_bottom(p, color_hex=GREEN_BG, size_pt=2)
    doc.add_paragraph()
    return p


def add_page_header(doc, icon: str, title: str, profile: str):
    """Adiciona cabeçalho de página com ícone e indicador de perfil colorido."""
    color_map = {
        "home/navegação":       "1A1A1A",   # preto
        "análise gerencial":    "FFBE00",   # amarelo
        "análise operacional":  "1A1A1A",   # preto
        "análise estratégica":  "FFBE00",   # amarelo
        "análise exploratória": "1A1A1A",   # preto
    }
    bg   = color_map.get(profile, GREEN_BG)
    tbl  = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"

    left = tbl.cell(0, 0)
    set_cell_bg(left, bg)
    left.width = Inches(0.15)
    left.text  = ""

    right = tbl.cell(0, 1)
    set_cell_bg(right, LIGHT_GRAY_BG)
    para  = right.paragraphs[0]
    run   = para.add_run(f"{icon}  {title}")
    run.font.bold      = True
    run.font.size      = Pt(13)
    run.font.color.rgb = JD_DARK
    doc.add_paragraph()


def add_info_box(doc, label: str, content, box_type="objective"):
    """Adiciona caixa informativa com borda colorida simulada via tabela."""
    border_color = {
        "objective": GREEN_BG,       # borda amarela
        "insight":   DARK_GREEN_BG,  # borda preta
        "note":      "808080",
        "kpi":       DARK_GREEN_BG,  # borda preta
    }.get(box_type, GREEN_BG)

    bg_color = {
        "objective": "FFFAE6",   # amarelo muito claro
        "insight":   "F5F5F5",   # cinza claro
        "note":      LIGHT_GRAY_BG,
        "kpi":       "F5F5F5",
    }.get(box_type, "FFFAE6")

    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_color)
    set_cell_border(cell, color_hex=border_color, width_pt=16)

    para = cell.paragraphs[0]
    if label:
        lbl_run = para.add_run(f"{label}  ")
        lbl_run.font.bold      = True
        lbl_run.font.size      = Pt(9.5)
        lbl_run.font.color.rgb = JD_DARK

    if isinstance(content, list):
        for i, item in enumerate(content):
            if i == 0:
                run = para.add_run(f"• {item}")
                run.font.size = Pt(9.5)
            else:
                np_ = cell.add_paragraph()
                r_  = np_.add_run(f"• {item}")
                r_.font.size = Pt(9.5)
    else:
        run = para.add_run(str(content))
        run.font.size = Pt(9.5)

    doc.add_paragraph()
    return tbl


def add_kpi_bullets(doc, items: list, icon="▸"):
    """Adiciona lista de bullets estilizada."""
    for item in items:
        p   = doc.add_paragraph()
        run = p.add_run(f"  {icon}  {item}")
        run.font.size = Pt(9.5)
        p.paragraph_format.left_indent  = Inches(0.2)
        p.paragraph_format.space_after  = Pt(2)


def add_styled_table(doc, headers: list, rows: list, header_bg=GREEN_BG, col_widths=None):
    """Adiciona tabela estilizada com cabeçalho e zebra striping."""
    if not rows:
        return None
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = "Table Grid"
    add_header_row(tbl, headers, bg_hex=header_bg)
    for i, row_data in enumerate(rows, 1):
        add_data_row(tbl, i, row_data, alternate=(i % 2 == 0))
    if col_widths:
        for i, w in enumerate(col_widths):
            if i < len(tbl.columns):
                for cell in tbl.columns[i].cells:
                    cell.width = Inches(w)
    doc.add_paragraph()
    return tbl


def add_separator(doc, color=GREEN_BG):
    """Adiciona linha separadora horizontal."""
    p = doc.add_paragraph()
    set_paragraph_border_bottom(p, color_hex=color, size_pt=1)
    p.paragraph_format.space_after = Pt(4)
    return p


def body_para(doc, text: str, size_pt=10, bold=False, color=None, indent=None):
    """Adiciona parágrafo de corpo de texto."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(4)
    return p


# ── Construtor principal do documento ─────────────────────────────────────────

def generate_docx(pbix_name: str, report: dict, analysis: dict, output_path: str,
                  llm_config=None, log_callback=None):
    """Gera o documento Word completo com 9 seções ricas de documentação."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

    meta          = analysis["metadata"]
    conn          = analysis["connections"]
    main_pages    = analysis["main_pages"]
    tooltip_pages = analysis["tooltip_pages"]
    page_cls      = analysis["page_classifications"]
    biz_entities  = analysis["biz_entities"]

    # ══════════════════════════════════════════════════════════════════════════
    # CAPA
    # ══════════════════════════════════════════════════════════════════════════

    # Faixa de topo verde escuro
    cover_tbl  = doc.add_table(rows=1, cols=1)
    cover_tbl.style = "Table Grid"
    cover_cell = cover_tbl.cell(0, 0)
    set_cell_bg(cover_cell, DARK_GREEN_BG)
    cover_para           = cover_cell.paragraphs[0]
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cover_para.add_run("  DOCUMENTAÇÃO DE RELATÓRIO POWER BI  ")
    cr.font.bold      = True
    cr.font.size      = Pt(11)
    cr.font.color.rgb = JD_YELLOW

    doc.add_paragraph()
    doc.add_paragraph()

    title_p           = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(pbix_name)
    tr.font.bold      = True
    tr.font.size      = Pt(28)
    tr.font.color.rgb = JD_GREEN

    doc.add_paragraph()

    sub_p           = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("Documentação Técnica e Analítica")
    sr.font.size      = Pt(18)
    sr.font.bold      = True
    sr.font.color.rgb = DARK_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    # Linha de acento amarela
    acc_tbl  = doc.add_table(rows=1, cols=1)
    acc_tbl.style = "Table Grid"
    acc_cell = acc_tbl.cell(0, 0)
    set_cell_bg(acc_cell, YELLOW_ACCENT)
    acc_p           = acc_cell.paragraphs[0]
    acc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    acc_r = acc_p.add_run("  Gerado automaticamente pelo PBIX Analyzer  ")
    acc_r.font.size      = Pt(9)
    acc_r.font.bold      = True
    acc_r.font.color.rgb = JD_DARK

    doc.add_paragraph()
    doc.add_paragraph()

    # Info de geração
    gen_p           = doc.add_paragraph()
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_p.add_run(f"Data de Geração: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    if meta.get("CreatedFromRelease"):
        rp           = doc.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp.add_run(f"Power BI Desktop Release: {meta['CreatedFromRelease']}")

    if conn.get("RemoteArtifacts"):
        ra           = conn["RemoteArtifacts"][0]
        cp           = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dsid = ra.get("DatasetId", "N/A")
        cp.add_run(f"Conectado ao Power BI Service  |  Dataset: {dsid[:20]}...")

    doc.add_paragraph()
    doc.add_paragraph()

    # Stats de capa (4 colunas)
    stats_tbl = doc.add_table(rows=1, cols=4)
    stats_tbl.style = "Table Grid"
    stats_data = [
        (str(analysis["total_pages"]),          "Páginas"),
        (str(analysis["total_visuals"]),         "Visuais"),
        (str(len(analysis["diagram_tables"])),   "Tabelas"),
        (str(len(analysis["all_fields"])),       "Campos"),
    ]
    for i, (val, lbl) in enumerate(stats_data):
        sc = stats_tbl.cell(0, i)
        set_cell_bg(sc, GREEN_BG)
        sp           = sc.paragraphs[0]
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = sp.add_run(val + "\n")
        vr.font.bold      = True
        vr.font.size      = Pt(20)
        vr.font.color.rgb = JD_DARK    # preto sobre amarelo
        lr = sp.add_run(lbl)
        lr.font.size      = Pt(9)
        lr.font.color.rgb = JD_DARK    # preto sobre amarelo

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # ÍNDICE
    # ══════════════════════════════════════════════════════════════════════════

    idx_hdr = doc.add_paragraph()
    idx_r   = idx_hdr.add_run("ÍNDICE")
    idx_r.font.bold      = True
    idx_r.font.size      = Pt(16)
    idx_r.font.color.rgb = JD_DARK    # preto (legível em fundo branco)
    set_paragraph_border_bottom(idx_hdr, color_hex=GREEN_BG, size_pt=3)
    doc.add_paragraph()

    toc_items = [
        ("1", "Resumo Executivo",                 "Visão geral do propósito, entidades e escopo do relatório"),
        ("2", "Visão Geral da Estrutura",          "Tabela de todas as páginas com perfil e composição de visuais"),
        ("3", "Arquitetura de Navegação",          "Como as páginas se conectam e o fluxo de uso recomendado"),
        ("4", "Páginas Principais",                "Análise detalhada de cada página principal do relatório"),
        ("5", "Tooltips e Páginas Auxiliares",     "Análise das páginas de tooltip e seu papel analítico"),
        ("6", "Inventário Técnico de Campos",      "Campos, medidas e dimensões catalogados por papel"),
        ("7", "Arquitetura do Modelo de Dados",    "Tabelas do modelo, tipos e papéis no relatório"),
        ("8", "Visuais Customizados",              "Custom visuals utilizados e suas funções analíticas"),
        ("9", "Informações Técnicas e Limitações", "Conexão, versão Power BI e notas técnicas"),
    ]

    toc_tbl = doc.add_table(rows=len(toc_items), cols=3)
    toc_tbl.style = "Table Grid"
    for i, (num, sec, desc) in enumerate(toc_items):
        bg = ALT_ROW_BG if i % 2 == 0 else "FFFFFF"
        c0, c1, c2 = toc_tbl.cell(i, 0), toc_tbl.cell(i, 1), toc_tbl.cell(i, 2)
        set_cell_bg(c0, GREEN_BG)
        set_cell_bg(c1, bg)
        set_cell_bg(c2, bg)
        c0.width = Inches(0.3)
        c1.width = Inches(2.8)
        c2.width = Inches(3.4)
        r0 = c0.paragraphs[0].add_run(num)
        r0.font.bold      = True
        r0.font.size      = Pt(10)
        r0.font.color.rgb = JD_DARK    # preto sobre amarelo
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = c1.paragraphs[0].add_run(sec)
        r1.font.bold      = True
        r1.font.size      = Pt(9.5)
        r1.font.color.rgb = JD_DARK
        r2 = c2.paragraphs[0].add_run(desc)
        r2.font.size      = Pt(9)
        r2.font.color.rgb = MID_GRAY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SEÇÃO 1 – RESUMO EXECUTIVO
    # ══════════════════════════════════════════════════════════════════════════

    add_section_header(doc, 1, "Resumo Executivo")

    body_para(doc,
        f"O relatório '{pbix_name}' é uma solução de Business Intelligence desenvolvida no Microsoft "
        f"Power BI para suporte à análise e tomada de decisão. O documento a seguir descreve sua "
        f"estrutura, conteúdo analítico e orientações de uso para diferentes perfis de usuário."
    )

    if biz_entities:
        body_para(doc,
            f"Os principais temas analíticos abordados são: {', '.join(biz_entities)}. "
            f"O relatório organiza essas perspectivas em {len(main_pages)} página(s) principal(is) "
            f"e {len(tooltip_pages)} página(s) de tooltip contextual."
        )

    # Perfis de usuário
    add_subsection_header(doc, "Perfis de Usuário e Páginas Recomendadas", level=3)

    profiles_found = {}
    for page in main_pages:
        prof = page_cls.get(page["name"], {}).get("profile", "análise estratégica")
        if prof not in profiles_found:
            profiles_found[prof] = []
        profiles_found[prof].append(page["name"])

    profile_descriptions = {
        "home/navegação":       "Ponto de entrada; recomendado para todos os usuários.",
        "análise gerencial":    "KPIs consolidados; recomendado para gestores e diretores.",
        "análise operacional":  "Detalhamento granular; recomendado para analistas e equipe operacional.",
        "análise estratégica":  "Posicionamento e estratégia; recomendado para líderes comerciais.",
        "análise exploratória": "Pesquisa e distribuição; recomendado para analistas de dados.",
    }

    profile_rows = []
    for prof, pages_list in profiles_found.items():
        profile_rows.append([
            prof.title(),
            ", ".join(pages_list),
            profile_descriptions.get(prof, "")
        ])
    if profile_rows:
        add_styled_table(doc,
            ["Perfil Analítico", "Páginas Recomendadas", "Descrição"],
            profile_rows,
            col_widths=[1.6, 2.2, 2.6]
        )

    if biz_entities:
        add_info_box(doc,
            "Entidades de Negócio Identificadas:",
            [e.title() for e in biz_entities],
            box_type="insight"
        )

    # Estatísticas rápidas
    add_subsection_header(doc, "Estatísticas do Relatório", level=3)
    stats_rows = [
        ["Total de Páginas",            str(analysis["total_pages"])],
        ["Páginas Principais",          str(len(main_pages))],
        ["Páginas de Tooltip",          str(len(tooltip_pages))],
        ["Total de Visuais",            str(analysis["total_visuals"])],
        ["Campos e Medidas Únicos",     str(len(analysis["all_fields"]))],
        ["Tabelas no Modelo",           str(len(analysis["diagram_tables"]))],
        ["Visuais Customizados em Uso", str(len(analysis["used_custom"]))],
    ]
    add_styled_table(doc, ["Indicador", "Valor"], stats_rows, col_widths=[3.5, 2.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SEÇÃO 2 – VISÃO GERAL DA ESTRUTURA
    # ══════════════════════════════════════════════════════════════════════════

    add_section_header(doc, 2, "Visão Geral da Estrutura")

    body_para(doc,
        "A tabela abaixo apresenta todas as páginas principais do relatório com seus respectivos "
        "perfis analíticos, composição de visuais e filtros de página disponíveis."
    )

    overview_rows = []
    for i, page in enumerate(main_pages, 1):
        cls         = page_cls.get(page["name"], {})
        content_c   = cls.get("content_count", 0)
        decor_c     = cls.get("decoration_count", 0)
        profile     = cls.get("profile", "—").title()
        filters_str = ", ".join(page["filters"][:3]) if page["filters"] else "—"
        if len(page["filters"]) > 3:
            filters_str += f" (+{len(page['filters'])-3})"
        overview_rows.append([
            str(i), page["name"], profile,
            f"{content_c} analíticos / {decor_c} decorativos",
            filters_str
        ])

    add_styled_table(doc,
        ["#", "Nome da Página", "Perfil", "Composição de Visuais", "Filtros de Página"],
        overview_rows,
        col_widths=[0.3, 1.7, 1.5, 1.9, 1.6]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SEÇÃO 3 – ARQUITETURA DE NAVEGAÇÃO
    # ══════════════════════════════════════════════════════════════════════════

    add_section_header(doc, 3, "Arquitetura de Navegação")

    home_pages  = [p for p in main_pages
                   if page_cls.get(p["name"], {}).get("profile") == "home/navegação"]
    other_mains = [p for p in main_pages
                   if page_cls.get(p["name"], {}).get("profile") != "home/navegação"]

    if home_pages:
        body_para(doc,
            f"O relatório utiliza a página '{home_pages[0]['name']}' como hub central de navegação. "
            f"A partir desta página, o usuário pode acessar as {len(other_mains)} seções analíticas "
            f"principais. O fluxo de navegação é intencional: da visão executiva às análises "
            f"operacionais detalhadas."
        )
    else:
        body_para(doc,
            f"O relatório é composto por {len(main_pages)} páginas principais organizadas em "
            f"sequência lógica analítica, do nível estratégico ao operacional."
        )

    add_subsection_header(doc, "Fluxo de Análise Recomendado", level=3)

    profile_order = [
        "home/navegação", "análise estratégica", "análise gerencial",
        "análise operacional", "análise exploratória"
    ]
    flow_items  = []
    used_pages  = set()
    for prof in profile_order:
        for page in main_pages:
            if page["name"] not in used_pages:
                if page_cls.get(page["name"], {}).get("profile") == prof:
                    flow_items.append(f"{page['name']} ({prof.title()})")
                    used_pages.add(page["name"])
    for page in main_pages:
        if page["name"] not in used_pages:
            flow_items.append(page["name"])

    add_kpi_bullets(doc, flow_items, icon="→")

    if tooltip_pages:
        doc.add_paragraph()
        add_info_box(doc,
            f"Padrão de Tooltips ({len(tooltip_pages)} páginas):",
            f"Este relatório utiliza {len(tooltip_pages)} páginas de tooltip para enriquecer a "
            f"experiência analítica. Ao passar o cursor sobre visuais específicos, contextos adicionais "
            f"são exibidos sem sair da página atual, reduzindo a necessidade de navegação.",
            box_type="insight"
        )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SEÇÃO 4 – PÁGINAS PRINCIPAIS
    # ══════════════════════════════════════════════════════════════════════════

    add_section_header(doc, 4, "Páginas Principais")

    PROFILE_ICONS = {
        "home/navegação":       "HOME",
        "análise gerencial":    "GERENCIAL",
        "análise operacional":  "OPERACIONAL",
        "análise estratégica":  "ESTRATEGICA",
        "análise exploratória": "EXPLORATORIA",
    }

    PROFILE_ICONS_EMOJI = {
        "home/navegação":       "🏠",
        "análise gerencial":    "📊",
        "análise operacional":  "🔍",
        "análise estratégica":  "📋",
        "análise exploratória": "🔬",
    }

    for page in main_pages:
        cls       = page_cls.get(page["name"], {})
        narrative = generate_page_narrative(page, cls)
        profile   = cls.get("profile", "análise estratégica")
        icon      = PROFILE_ICONS_EMOJI.get(profile, "📄")
        deep      = narrative.get("deep", {})

        # Enriquecimento opcional via LLM — silently falls back to heuristics on any failure
        if llm_config and llm_config.get("provider") not in (None, "Desabilitado"):
            if log_callback:
                log_callback(f"🤖 Gerando narrativa LLM para página '{page['name']}'...")
            llm_result = enhance_page_with_llm(page, deep, cls, llm_config)
            if llm_result:
                if llm_result.get("objetivo"):
                    narrative["objective"] = llm_result["objetivo"]
                if llm_result.get("resumo"):
                    narrative["executive_summary"] = llm_result["resumo"]
                if llm_result.get("leitura"):
                    narrative["reading_guide"] = llm_result["leitura"]

        add_page_header(doc, icon, page["name"], profile)

        # Objetivo
        add_info_box(doc, "Objetivo:", narrative["objective"], box_type="objective")

        # Resumo executivo (texto narrativo rico)
        body_para(doc, narrative["executive_summary"])

        # Destaques (key insights)
        key_ins = cls.get("key_insights", [])
        if key_ins:
            add_info_box(doc, "Destaques Analíticos:", key_ins, box_type="insight")

        # ── Composição visual ────────────────────────────────────────────
        add_subsection_header(doc, "Composição dos Visuais", level=3)
        visuals   = page["visuals"]
        content_v = [v for v in visuals if v["type"] not in DECORATION_TYPES]
        deco_v    = [v for v in visuals if v["type"] in DECORATION_TYPES]
        vc_count  = defaultdict(int)
        for v in content_v:
            vc_count[v["label"]] += 1
        vc_rows = [(lbl, str(cnt)) for lbl, cnt in sorted(vc_count.items(), key=lambda x: -x[1])]
        if vc_rows:
            add_styled_table(doc, ["Tipo de Visual Analítico", "Qtd."], vc_rows,
                             col_widths=[3.5, 1.0])
        body_para(doc,
            f"Elementos decorativos/layout (formas, textos, botões, imagens): {len(deco_v)}.",
            size_pt=8.5, color=MID_GRAY)

        # ── KPIs: usa os nomes reais dos cards ────────────────────────────
        kpis_deep = deep.get("kpis", [])
        if kpis_deep:
            add_subsection_header(doc, "Principais Indicadores (KPIs)", level=3)
            add_info_box(doc, "", kpis_deep[:12], box_type="kpi")

        # ── Gráficos: descreve o que cada um analisa ──────────────────────
        chart_details = deep.get("chart_details", [])
        if chart_details:
            add_subsection_header(doc, "O Que os Gráficos Analisam", level=3)
            chart_rows = []
            for ch in chart_details:
                meas_str = ", ".join(ch["measures"]) if ch["measures"] else "—"
                dims_str = ", ".join(ch["dims"])     if ch["dims"]     else "—"
                chart_rows.append([ch["type"], meas_str, dims_str])
            add_styled_table(doc,
                ["Tipo de Gráfico", "Medidas / Valores", "Dimensões / Eixo"],
                chart_rows, col_widths=[1.8, 2.2, 2.2])

        # ── Tabelas/Matrizes: colunas reais ───────────────────────────────
        table_cols = deep.get("table_cols", [])
        if table_cols:
            add_subsection_header(doc, "Detalhamento em Tabelas e Matrizes", level=3)
            for i, cols in enumerate(table_cols[:3], 1):
                col_str = " | ".join(cols[:8])
                extra   = f"  (+{len(cols)-8} colunas)" if len(cols) > 8 else ""
                body_para(doc, f"Tabela {i}: {col_str}{extra}", size_pt=9)

        # ── Análise estatística ───────────────────────────────────────────
        stat_details = deep.get("stat_details", [])
        if stat_details:
            add_subsection_header(doc, "Análise Estatística", level=3)
            for st in stat_details[:2]:
                var_str = ", ".join(st["vars"][:5])
                body_para(doc, f"{st['type']} — variáveis: {var_str}", size_pt=9)

        # ── Filtros interativos: campos reais dos slicers ─────────────────
        filters_deep = deep.get("filters", [])
        map_details  = deep.get("map_details", [])
        if filters_deep or map_details:
            add_subsection_header(doc, "Dimensões e Filtros Disponíveis", level=3)
            filter_items = [f"🔘 {f}" for f in filters_deep]
            if map_details:
                filter_items += [f"🗺️  Geo: {mp['geo']}  →  {mp['measure']}"
                                 for mp in map_details]
            add_info_box(doc, "", filter_items[:16], box_type="note")

        # ── Tabelas de dados do modelo ────────────────────────────────────
        page_tables = sorted({t for v in visuals for t in v["tables"] if t})
        if page_tables:
            add_subsection_header(doc, "Tabelas de Dados Utilizadas", level=3)
            tbl_items = []
            for tn in page_tables:
                tc = analysis["table_classifications"].get(tn, {})
                tbl_items.append(f"{tn}  ({tc.get('type', '—')})")
            add_kpi_bullets(doc, tbl_items, icon="▸")

        # ── Guia de leitura ───────────────────────────────────────────────
        add_subsection_header(doc, "Como Interpretar Esta Página", level=3)
        body_para(doc, narrative["reading_guide"])

        # Callouts especiais
        callouts = []
        if cls.get("has_maps"):
            callouts.append("Mapa geográfico disponível para análise de distribuição regional.")
        if cls.get("has_stats_visuals"):
            callouts.append(
                "Gráficos Box & Whisker e/ou Dispersão permitem análise estatística de distribuição e outliers."
            )
        if cls.get("has_matrices"):
            callouts.append("Matrizes/tabelas suportam drill-down para análise no nível de detalhe desejado.")
        if cls.get("has_many_slicers"):
            callouts.append(
                f"Com {cls['slicer_count']} filtros interativos, esta página oferece alta flexibilidade analítica."
            )
        if callouts:
            add_info_box(doc, "Recursos Analíticos Especiais:", callouts, box_type="insight")

        add_separator(doc)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SEÇÃO 5 – TOOLTIPS E PÁGINAS AUXILIARES
    # ══════════════════════════════════════════════════════════════════════════

    add_section_header(doc, 5, "Tooltips e Páginas Auxiliares")

    body_para(doc,
        "Tooltips são páginas especiais do Power BI que aparecem em sobreposição quando o usuário "
        "posiciona o cursor sobre um visual em outra página. Elas fornecem profundidade analítica "
        "contextual sem exigir navegação, mantendo o usuário focado na análise principal."
    )

    if tooltip_pages:
        body_para(doc,
            f"Este relatório conta com {len(tooltip_pages)} páginas de tooltip, demonstrando alto "
            f"grau de sofisticação na experiência analítica. Cada tooltip foi projetado para enriquecer "
            f"visuais específicos com informações complementares relevantes."
        )
        doc.add_paragraph()

        add_subsection_header(doc, "Inventário de Tooltips", level=3)

        tt_rows = []
        for page in tooltip_pages:
            content_v = [v for v in page["visuals"] if v["type"] not in DECORATION_TYPES]
            flds = []
            for v in content_v:
                for f in v["fields"]:
                    if f["display"] and f["display"] not in flds:
                        flds.append(f["display"])
            vtype_labels = list(set(v["label"] for v in content_v))[:3]
            tipo_info = ", ".join(vtype_labels) or "—"
            flds_str  = ", ".join(flds[:5])
            if len(flds) > 5:
                flds_str += f" (+{len(flds)-5})"
            tt_rows.append([page["name"], str(len(page["visuals"])), tipo_info, flds_str or "—"])

        add_styled_table(doc,
            ["Nome do Tooltip", "Visuais", "Tipos de Visual", "Campos / Métricas"],
            tt_rows,
            col_widths=[1.7, 0.6, 1.8, 2.9]
        )

        # Análise por grupo temático
        price_tt  = [p for p in tooltip_pages
                     if any(w in p["name"].lower() for w in ["preco", "preço", "price", "apn"])]
        stats_tt  = [p for p in tooltip_pages
                     if any(w in p["name"].lower() for w in ["distribu", "estat", "box", "whisker"])]
        model_tt  = [p for p in tooltip_pages
                     if any(w in p["name"].lower() for w in ["modelo", "familia", "produto"])]
        dealer_tt = [p for p in tooltip_pages
                     if any(w in p["name"].lower() for w in ["dealer", "canal", "concession"])]

        narrative_tt = []
        if price_tt:
            narrative_tt.append(
                f"Os {len(price_tt)} tooltip(s) de preço fornecem detalhamento de precificação "
                f"contextual diretamente nos visuais de análise."
            )
        if stats_tt:
            narrative_tt.append(
                f"Os {len(stats_tt)} tooltip(s) estatístico(s) apresentam distribuições e "
                f"análises de variância no nível de detalhe."
            )
        if model_tt:
            narrative_tt.append(
                f"Os {len(model_tt)} tooltip(s) de modelo/família detalham a composição "
                f"do portfólio por família ou linha de produto."
            )
        if dealer_tt:
            narrative_tt.append(
                f"Os {len(dealer_tt)} tooltip(s) de dealer/canal fornecem contexto de "
                f"desempenho individual de canais de venda."
            )
        if not narrative_tt:
            narrative_tt.append(
                f"Os {len(tooltip_pages)} tooltips enriquecem a experiência analítica com "
                f"informações contextuais adicionais ao nível de detalhe dos visuais principais."
            )

        add_info_box(doc, "Análise dos Grupos de Tooltip:", narrative_tt, box_type="insight")
    else:
        body_para(doc, "Nenhuma página de tooltip identificada neste relatório.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SEÇÃO 6 – INVENTÁRIO TÉCNICO DE CAMPOS
    # ══════════════════════════════════════════════════════════════════════════

    add_section_header(doc, 6, "Inventário Técnico de Campos")

    body_para(doc,
        "Esta seção cataloga todos os campos, medidas e colunas identificados no relatório, "
        "organizados por relevância e papel analítico. Campos de alta relevância são aqueles "
        "presentes em duas ou mais páginas do relatório."
    )

    sorted_fields = sorted(analysis["all_fields"].items(), key=lambda x: -x[1]["count"])

    # Alta relevância
    high_rel = [(k, v) for k, v in sorted_fields if v["count"] >= 2]
    if high_rel:
        add_subsection_header(doc, "Campos de Alta Relevância (presentes em 2+ páginas)", level=3)
        hr_rows = []
        for fname, finfo in high_rel:
            role = classify_field(fname, finfo["kind"])
            ps   = ", ".join(sorted(finfo["pages"]))
            if len(ps) > 60:
                ps = ps[:57] + "..."
            hr_rows.append([fname, finfo.get("table", "—"), role, finfo["kind"], ps])
        add_styled_table(doc,
            ["Campo / Medida", "Tabela", "Papel Analítico", "Tipo", "Páginas"],
            hr_rows,
            header_bg="B88600",
            col_widths=[1.5, 1.2, 1.3, 0.8, 1.7]
        )

    # KPIs / Medidas
    measures = [(k, v) for k, v in sorted_fields if v["kind"] == "Medida"]
    if measures:
        add_subsection_header(doc, "KPIs e Medidas", level=3)
        m_rows = []
        for fname, finfo in measures:
            ps = ", ".join(sorted(finfo["pages"]))
            if len(ps) > 50:
                ps = ps[:47] + "..."
            m_rows.append([fname, finfo.get("table", "—"), str(finfo["count"]), ps])
        add_styled_table(doc,
            ["Nome da Medida", "Tabela de Origem", "Usos", "Páginas que Utilizam"],
            m_rows,
            col_widths=[1.8, 1.4, 0.5, 2.8]
        )

    # Dimensões e colunas por papel
    columns = [(k, v) for k, v in sorted_fields if v["kind"] != "Medida"]
    if columns:
        add_subsection_header(doc, "Dimensões e Colunas", level=3)
        role_groups = defaultdict(list)
        for fname, finfo in columns:
            role = classify_field(fname, finfo["kind"])
            role_groups[role].append((fname, finfo))

        for role_name, role_fields in sorted(role_groups.items()):
            if not role_fields:
                continue
            body_para(doc, f"  {role_name}:", size_pt=9.5, bold=True, color=JD_DARK)
            rf_rows = []
            for rfname, rfinfo in sorted(role_fields, key=lambda x: -x[1]["count"]):
                ps = ", ".join(sorted(rfinfo["pages"]))
                if len(ps) > 50:
                    ps = ps[:47] + "..."
                rf_rows.append([rfname, rfinfo.get("table", "—"), rfinfo["kind"], ps])
            add_styled_table(doc,
                ["Nome", "Tabela de Origem", "Tipo", "Páginas que Utilizam"],
                rf_rows,
                col_widths=[1.8, 1.5, 0.8, 2.4]
            )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SEÇÃO 7 – ARQUITETURA DO MODELO DE DADOS
    # ══════════════════════════════════════════════════════════════════════════

    add_section_header(doc, 7, "Arquitetura do Modelo de Dados")

    body_para(doc,
        "O modelo de dados do relatório é composto pelas tabelas listadas abaixo, cada uma com "
        "papel específico na estrutura analítica. As tabelas foram classificadas automaticamente "
        "com base em convenções de nomenclatura e padrões de uso no relatório."
    )

    # Resumo de tipos
    type_counts = defaultdict(int)
    for tc in analysis["table_classifications"].values():
        type_counts[tc["type"]] += 1
    if type_counts:
        add_info_box(doc, "Composição do Modelo:",
            [f"{t.title()}: {c} tabela(s)" for t, c in sorted(type_counts.items())],
            box_type="objective"
        )

    add_subsection_header(doc, "Inventário de Tabelas", level=3)
    tbl_rows = []
    for tname in analysis["all_tables_set"]:
        tc        = analysis["table_classifications"].get(tname, {"type": "—", "role": "—"})
        usage     = analysis["all_tables_used"].get(tname, {})
        ps        = ", ".join(sorted(usage.get("pages", set()))) or "—"
        if len(ps) > 50:
            ps = ps[:47] + "..."
        in_diag = "Sim" if tname in analysis["diagram_tables"] else "—"
        tbl_rows.append([tname, tc["type"].title(), tc["role"][:58], ps, in_diag])

    add_styled_table(doc,
        ["Nome da Tabela", "Tipo", "Papel no Relatório", "Páginas que Utilizam", "Diagrama"],
        tbl_rows,
        col_widths=[1.3, 0.9, 2.2, 1.9, 0.7]
    )

    # Narrativa sobre centralidade do modelo
    fato_t = [n for n, tc in analysis["table_classifications"].items() if tc["type"] == "fato"]
    dim_t  = [n for n, tc in analysis["table_classifications"].items() if tc["type"] == "dimensão"]
    cal_t  = [n for n, tc in analysis["table_classifications"].items() if tc["type"] == "calendário"]

    parts = []
    if fato_t:
        parts.append(
            f"A(s) tabela(s) de fatos ({', '.join(fato_t)}) "
            f"{'forma' if len(fato_t)==1 else 'formam'} o núcleo analítico do modelo, "
            f"contendo os dados transacionais ou de pesquisa."
        )
    if dim_t:
        dim_str = ', '.join(dim_t[:4]) + ('...' if len(dim_t) > 4 else '')
        parts.append(
            f"As {len(dim_t)} tabela(s) de dimensão ({dim_str}) fornecem o contexto de "
            f"filtragem e agrupamento para as análises."
        )
    if cal_t:
        parts.append(
            f"A tabela de calendário ({', '.join(cal_t)}) habilita análises temporais "
            f"com filtros por data, mês, trimestre e ano."
        )
    if parts:
        body_para(doc, " ".join(parts))

    # ── Seção de Relacionamentos ────────────────────────────────────────────
    relationships = analysis.get("relationships", [])
    if relationships:
        add_subsection_header(doc, "Relacionamentos entre Tabelas (Inferidos)", level=3)

        body_para(doc,
            "Os relacionamentos abaixo foram identificados de forma automática a partir da "
            "análise de co-uso de tabelas nos visuais do relatório. A confiança reflete a "
            "frequência com que as tabelas aparecem juntas e a presença de colunas de junção "
            "com nomes similares. Este mapeamento substitui de forma aproximada o diagrama "
            "oficial do modelo, cujos dados estão armazenados em formato binário no arquivo."
        )

        rel_rows = []
        for r in relationships:
            via_txt = r["via"][:45] if r["via"] else "—"
            pages_txt = ", ".join(r["pages"][:2]) + ("..." if len(r["pages"]) > 2 else "")
            rel_rows.append([r["from_table"], r["to_table"], via_txt,
                             r["confidence"], str(r["count"]), pages_txt])

        add_styled_table(doc,
            ["Tabela A", "Tabela B", "Via / Coluna Provável",
             "Confiança", "Co-usos", "Páginas"],
            rel_rows,
            col_widths=[1.3, 1.3, 1.8, 0.8, 0.6, 1.2]
        )

        # Diagrama visual embutido
        img_bytes = generate_relationship_diagram(
            relationships,
            analysis["all_tables_set"],
            analysis["table_classifications"]
        )
        if img_bytes:
            import io as _io
            add_subsection_header(doc, "Diagrama de Relacionamentos", level=3)
            body_para(doc,
                "O diagrama abaixo representa graficamente os relacionamentos inferidos entre "
                "tabelas. Cada nó corresponde a uma tabela do modelo; a espessura e estilo das "
                "linhas indicam o nível de confiança do relacionamento detectado."
            )
            try:
                doc.add_picture(_io.BytesIO(img_bytes), width=Inches(6.5))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SEÇÃO 8 – VISUAIS CUSTOMIZADOS
    # ══════════════════════════════════════════════════════════════════════════

    add_section_header(doc, 8, "Visuais Customizados")

    body_para(doc,
        "Visuais customizados (Custom Visuals) são componentes desenvolvidos pela comunidade ou "
        "por fornecedores certificados que estendem as capacidades nativas do Power BI. Sua "
        "utilização neste relatório indica requisitos analíticos que os visuais nativos não "
        "atendiam plenamente."
    )

    used_cv   = analysis["used_custom"]
    unused_cv = analysis["unused_custom"]

    if used_cv:
        add_subsection_header(doc, f"Visuais Customizados em Uso ({len(used_cv)})", level=3)
        for cv_key, cv_info in used_cv.items():
            use_count = analysis["visual_type_count"].get(cv_key, 0)
            cv_tbl  = doc.add_table(rows=1, cols=1)
            cv_tbl.style = "Table Grid"
            cv_cell = cv_tbl.cell(0, 0)
            set_cell_bg(cv_cell, "FFF8DC")   # bege amarelado leve
            set_cell_border(cv_cell, color_hex=GREEN_BG, width_pt=12)
            p0  = cv_cell.paragraphs[0]
            r0  = p0.add_run(f"  {cv_info['name']}  ")
            r0.font.bold      = True
            r0.font.size      = Pt(10.5)
            r0.font.color.rgb = JD_DARK
            rc  = p0.add_run(f"({use_count} uso{'s' if use_count != 1 else ''} no relatório)")
            rc.font.size      = Pt(9)
            rc.font.color.rgb = MID_GRAY
            p1  = cv_cell.add_paragraph()
            rd  = p1.add_run(f"  {cv_info['desc']}")
            rd.font.size = Pt(9.5)
            p2  = cv_cell.add_paragraph()
            bld = p2.add_run("  Valor analítico: ")
            bld.font.bold = True
            rv  = p2.add_run(cv_info["value"])
            rv.font.size      = Pt(9.5)
            rv.font.color.rgb = JD_DARK
            doc.add_paragraph()

    if unused_cv:
        add_subsection_header(doc,
            f"Visuais Disponíveis mas Não Utilizados ({len(unused_cv)})", level=3
        )
        body_para(doc,
            "Os visuais abaixo estão incluídos no pacote do relatório mas não foram encontrados "
            "em nenhuma página. Podem ter sido incluídos para uso futuro ou removidos durante "
            "revisões anteriores.",
            color=MID_GRAY
        )
        unused_rows = [[info["name"], info["desc"][:80]] for _, info in unused_cv.items()]
        add_styled_table(doc, ["Visual", "Descrição"], unused_rows,
                         header_bg="808080", col_widths=[2.0, 4.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SEÇÃO 9 – INFORMAÇÕES TÉCNICAS E LIMITAÇÕES
    # ══════════════════════════════════════════════════════════════════════════

    add_section_header(doc, 9, "Informações Técnicas e Limitações")

    add_subsection_header(doc, "Conexão com Dataset", level=3)
    if conn.get("RemoteArtifacts"):
        ra = conn["RemoteArtifacts"][0]
        tech_rows = [
            ["Tipo de Conexão", "Dataset Remoto (Power BI Service)"],
            ["Dataset ID",      ra.get("DatasetId", "N/A")],
            ["Report ID",       ra.get("ReportId", "N/A")],
            ["Workspace ID",    conn.get("OriginalWorkspaceObjectId", "N/A")],
        ]
        add_info_box(doc, "Nota sobre Dataset Remoto:",
            "Este relatório está conectado a um dataset remoto no Power BI Service. "
            "As definições de medidas DAX não estão disponíveis no arquivo .pbix local. "
            "Para visualizar as fórmulas DAX completas, acesse o dataset no Power BI Service.",
            box_type="note"
        )
    else:
        tech_rows = [["Tipo de Conexão", "Dataset Local / Importação Direta"]]

    if meta:
        tech_rows += [
            ["Versão do Metadata",  str(meta.get("Version", "N/A"))],
            ["Power BI Release",    meta.get("CreatedFromRelease", "N/A")],
            ["Criado a partir de",  meta.get("CreatedFrom", "N/A")],
        ]

    add_styled_table(doc, ["Parâmetro Técnico", "Valor"], tech_rows,
                     col_widths=[2.5, 4.0])

    # Distribuição de visuais
    add_subsection_header(doc, "Distribuição de Tipos de Visuais", level=3)
    vt_items  = [(get_visual_label(k), v) for k, v in
                 sorted(analysis["visual_type_count"].items(), key=lambda x: -x[1])]
    total_vt  = sum(v for _, v in vt_items) or 1
    vt_rows   = [(lbl, str(cnt), f"{cnt/total_vt*100:.1f}%") for lbl, cnt in vt_items]
    add_styled_table(doc, ["Tipo de Visual", "Quantidade", "% do Total"], vt_rows,
                     col_widths=[3.5, 1.0, 1.0])

    # Limitações
    add_subsection_header(doc, "Limitações e Notas", level=3)
    limitations = [
        "As definições de fórmulas DAX das medidas não estão disponíveis quando o modelo "
        "está conectado a um dataset remoto no Power BI Service.",
        "A classificação de páginas, tabelas e campos é baseada em heurísticas de nomenclatura "
        "e pode não refletir 100% da intenção original do desenvolvedor.",
        "Visuais com configuração dinâmica (via parâmetros ou bookmarks) podem não ter todos "
        "os campos detectados estaticamente.",
        "Esta documentação foi gerada automaticamente em "
        + datetime.now().strftime("%d/%m/%Y às %H:%M") + ".",
    ]
    add_info_box(doc, "Limitações Conhecidas:", limitations, box_type="note")

    # Rodapé
    doc.add_paragraph()
    footer_p           = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "Documento gerado automaticamente pelo PBIX Analyzer  |  "
        "Para suporte consulte o repositório da ferramenta."
    )
    fr.font.size      = Pt(8)
    fr.font.color.rgb = MID_GRAY
    fr.font.italic    = True

    doc.save(output_path)
    return output_path


# ==============================================================================
# PONTO DE ENTRADA
# ==============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Analisa um arquivo .pbix e gera documentação rica em .docx"
    )
    parser.add_argument("pbix",           help="Caminho para o arquivo .pbix")
    parser.add_argument("--output", "-o", help="Caminho de saída para o .docx (opcional)")
    args = parser.parse_args()

    pbix_path = args.pbix
    if not os.path.exists(pbix_path):
        print(f"ERRO: Arquivo não encontrado: {pbix_path}")
        sys.exit(1)

    pbix_name = Path(pbix_path).stem
    if args.output:
        output_path = args.output
    else:
        output_dir  = Path(pbix_path).parent
        output_path = str(output_dir / f"{pbix_name}_Documentacao.docx")

    print(f"Lendo: {pbix_path}")
    print("Extraindo arquivo .pbix...")
    extract_dir = extract_pbix(pbix_path)

    try:
        print("Parseando Layout do Relatorio...")
        report = parse_layout(extract_dir)

        print("Parseando tabelas do modelo...")
        tables = parse_diagram_layout(extract_dir)

        print("Lendo informacoes de conexao...")
        connections = parse_connections(extract_dir)
        metadata    = parse_metadata(extract_dir)

        print("Analisando e classificando estrutura...")
        analysis = build_analysis(report, tables, connections, metadata)

        n_main    = len(analysis["main_pages"])
        n_tooltip = len(analysis["tooltip_pages"])
        print(f"   -> {analysis['total_pages']} paginas ({n_main} principais, {n_tooltip} tooltips)")
        print(f"   -> {analysis['total_visuals']} visuais | {len(analysis['all_fields'])} campos | {len(tables)} tabelas")
        entities = ', '.join(analysis['biz_entities'])
        print(f"   -> Entidades detectadas: {entities}")

        print("Gerando documentacao Word (9 secoes)...")
        generate_docx(pbix_name, report, analysis, output_path)

        print(f"\nDocumentacao gerada com sucesso!")
        print(f"   Arquivo: {output_path}")

    finally:
        shutil.rmtree(extract_dir, ignore_errors=True)


if __name__ == "__main__":
    main()
