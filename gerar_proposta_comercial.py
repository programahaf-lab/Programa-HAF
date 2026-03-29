from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ModuleNotFoundError as exc:
    raise SystemExit(
        "Dependencia ausente: instale 'python-docx' antes de executar.\n"
        "Exemplo: pip install python-docx"
    ) from exc


# ============================================================================
# CAMPOS EDITAVEIS DO TEMPLATE
# ============================================================================
CLIENTE = "Empresa Cliente"
PROJETO = "Documentacao de Artefatos Analiticos"
DATA_PROPOSTA = date.today().strftime("%d/%m/%Y")
NUMERO_PROPOSTA = "PC-2026-001"
PACOTE_SELECIONADO = "Pacote 3 | Executivo"
INVESTIMENTO_FINAL = "R$ 1.250,00"
SEU_NOME = "Seu Nome"
NOME_MARCA = "Sua Marca"
WHATSAPP = "+55 11 99999-9999"
EMAIL = "contato@suamarca.com"
LINKEDIN = "linkedin.com/in/seuperfil"
ARQUIVO_SAIDA = "proposta_comercial_template.docx"


def criar_configuracao_proposta(
    cliente=CLIENTE,
    projeto=PROJETO,
    data_proposta=DATA_PROPOSTA,
    numero_proposta=NUMERO_PROPOSTA,
    pacote_selecionado=PACOTE_SELECIONADO,
    investimento_final=INVESTIMENTO_FINAL,
    nome_saida=ARQUIVO_SAIDA,
):
    """Centraliza os campos variaveis da proposta para uso por script ou integracao futura."""
    return {
        "cliente": cliente,
        "projeto": projeto,
        "data_proposta": data_proposta,
        "numero_proposta": numero_proposta,
        "pacote_selecionado": pacote_selecionado,
        "investimento_final": investimento_final,
        "nome_saida": nome_saida,
        "seu_nome": SEU_NOME,
        "nome_marca": NOME_MARCA,
        "whatsapp": WHATSAPP,
        "email": EMAIL,
        "linkedin": LINKEDIN,
    }


# ============================================================================
# CONFIGURACOES VISUAIS
# ============================================================================
COR_PRIMARIA = RGBColor(22, 28, 36)
COR_SECUNDARIA = RGBColor(88, 99, 112)
COR_DESTAQUE = RGBColor(12, 61, 89)
COR_BORDA = "D6DEE6"
COR_TABELA_HEADER = "16242F"
COR_TABELA_ZEBRA = "F7F9FB"
COR_TABELA_SELECIONADA = "E7F0F6"
COR_BLOCO = "F4F7F9"
COR_CAPA = "EEF3F6"
COR_MUTED = RGBColor(108, 117, 125)
FONTE_PADRAO = "Calibri"


RESUMO_EXECUTIVO = (
    "Esta proposta contempla a producao de documentacao analitica profissional "
    "para notebooks, dashboards e artefatos de dados, com foco em transformar "
    "entregas tecnicas em materiais claros, valorizados e prontos para uso "
    "corporativo.\n\n"
    "O trabalho combina leitura tecnica do artefato, organizacao do fluxo "
    "operacional, inventario dos componentes relevantes e, quando aplicavel, "
    "uma camada executiva pronta para compartilhamento com lideranca, cliente "
    "ou areas de negocio."
)

CONTEXTO_INTRO = (
    "Em muitos projetos de dados, o artefato final existe, mas o conhecimento "
    "fica disperso no codigo, no dashboard ou na explicacao verbal. Isso gera "
    "dificuldade para:"
)

CONTEXTO_LISTA = [
    "entender rapidamente como o artefato esta organizado",
    "identificar fontes, componentes, consultas, funcoes ou visuais relevantes",
    "repassar conhecimento entre times tecnicos, operacionais e gestores",
    "sustentar a solucao ao longo do tempo com menor dependencia de contexto informal",
    "comunicar o valor da entrega para publico nao tecnico com mais clareza",
]

CONTEXTO_FINAL = (
    "A proposta deste servico e elevar a percepcao de valor da entrega, "
    "convertendo o artefato analisado em um material documental com padrao "
    "profissional, rastreabilidade tecnica e leitura executiva adequada para "
    "uso interno ou apresentacao ao cliente."
)

OBJETIVOS = [
    "documentar o artefato com estrutura formal e leitura objetiva",
    "registrar fluxo, dependencias, regras e componentes tecnicos relevantes",
    "produzir um inventario tecnico utilizavel por times de dados e negocio",
    "reduzir dependencia de explicacoes manuais para entendimento da entrega",
    "aumentar a clareza, a percepcao de organizacao e o valor percebido da solucao",
]

ESCOPO_TECNICO = [
    "leitura tecnica do arquivo fonte enviado",
    "interpretacao da estrutura do notebook, dashboard ou artefato analitico",
    "mapeamento do fluxo logico, etapas, consultas, transformacoes, metricas e dependencias",
    "inventario dos componentes relevantes do artefato analisado",
    "organizacao da documentacao em formato tecnico estruturado e rastreavel",
]

ESCOPO_EXECUTIVO = [
    "consolidacao de uma visao executiva do projeto ou relatorio",
    "traducao do conteudo tecnico para linguagem acessivel ao publico de gestao",
    "sintese do valor analitico, uso do artefato e leitura recomendada",
    "criacao de material executivo em HTML quando previsto no pacote contratado",
]

ESCOPO_REVISAO = [
    "revisao final da entrega",
    "ajustes finos dentro da quantidade de revisoes prevista no pacote contratado",
]

ENTREGAVEIS = [
    "documentacao tecnica estruturada do artefato analisado",
    "resumo executivo com leitura objetiva da entrega",
    "inventario tecnico com os principais componentes identificados",
    "registro do fluxo operacional do artefato",
    "apresentacao executiva em HTML quando prevista no pacote",
]

NAO_INCLUIDOS = [
    "desenvolvimento do notebook ou dashboard",
    "correcao de bugs ou refatoracao do projeto",
    "alteracao de regras de negocio ja implementadas",
    "criacao de novas analises alem do material recebido",
    "participacao em reunioes extras nao previstas",
    "sustentacao continua ou suporte recorrente apos entrega",
    "auditoria completa do ambiente, dados ou arquitetura da solucao",
]

METODOLOGIA = [
    ("Recebimento do material", "Coleta dos arquivos, contexto inicial e confirmacao do tipo de artefato a ser analisado."),
    ("Leitura estrutural", "Analise do artefato para identificar composicao, fluxo, dependencias, regras e pontos de atencao."),
    ("Inventario tecnico", "Catalogacao dos componentes relevantes em formato documental claro e reutilizavel."),
    ("Consolidacao executiva", "Sintese do conteudo para leitura gerencial, apresentacao ou repasse interno quando aplicavel."),
    ("Entrega e ajustes", "Compartilhamento do material final e refinamentos previstos dentro do pacote contratado."),
]

PRAZOS = [
    "artefatos simples: 1 a 2 dias uteis",
    "artefatos com maior densidade tecnica: 2 a 4 dias uteis",
    "projetos com multiplos arquivos ou maior profundidade: prazo sob analise",
]

PACOTES = [
    [
        "Pacote 1",
        "Essencial Tecnico",
        "Documentacao tecnica objetiva para artefatos menores, com foco em estrutura, fluxo e componentes principais",
        "R$ 290 a R$ 450",
    ],
    [
        "Pacote 2",
        "Tecnico Avancado",
        "Documentacao tecnica aprofundada com inventario estrutural, dependencias e maior nivel de rastreabilidade",
        "R$ 590 a R$ 900",
    ],
    [
        "Pacote 3",
        "Executivo",
        "Documentacao tecnica estruturada + camada executiva para leitura gerencial, apresentacao comercial ou uso com cliente final",
        "R$ 890 a R$ 1.500",
    ],
    [
        "Pacote 4",
        "Projeto Sob Medida",
        "Projetos com multiplos artefatos, maior profundidade tecnica ou necessidade de organizacao documental ampliada",
        "A partir de R$ 1.290",
    ],
]

CONDICOES_COMERCIAIS = [
    "pagamento via PIX ou outro meio previamente acordado",
    "sugestao de condicao: 50% no aceite e 50% na entrega",
    "a proposta tem validade de 7 dias corridos",
    "revisoes adicionais fora do pacote poderao ser orcadas separadamente",
    "demandas extras ou alteracoes de escopo poderao impactar prazo e investimento",
]

PREMISSAS = [
    "os arquivos serao enviados de forma integra e acessivel",
    "o cliente fornecera contexto minimo quando necessario",
    "o servico sera realizado com base nos artefatos disponibilizados",
    "eventuais informacoes ausentes podem limitar a profundidade da documentacao",
    "o material produzido tem finalidade documental e executiva, nao substituindo governanca interna, homologacao tecnica ou aprovacao formal de arquitetura",
]

DIFERENCIAIS = [
    "estrutura documental com padrao profissional",
    "inventario tecnico do artefato analisado",
    "rastreabilidade de fluxo, componentes e dependencias",
    "clareza para publicos tecnicos e nao tecnicos",
    "padronizacao visual e organizacional da entrega",
    "material pronto para compartilhamento interno, apresentacao externa ou repasse entre times",
]

PROXIMOS_PASSOS = [
    "aprovacao da proposta",
    "envio dos arquivos e contexto necessario",
    "confirmacao do pagamento inicial",
    "inicio da execucao",
    "entrega do material dentro do prazo acordado",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=COR_BORDA, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_border(paragraph, color=COR_BORDA, size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    border = p_pr.find(qn("w:pBdr"))
    if border is None:
        border = OxmlElement("w:pBdr")
        p_pr.append(border)

    bottom = border.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        border.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def configure_page(section):
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def apply_base_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = FONTE_PADRAO
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = COR_PRIMARIA
    pf = normal.paragraph_format
    pf.space_after = Pt(5)
    pf.line_spacing = 1.2


def format_run(run, size=None, bold=False, color=None, font_name=FONTE_PADRAO, all_caps=False):
    run.font.name = font_name
    run.font.bold = bold
    run.font.all_caps = all_caps
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_title(document, text, subtitle=False):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8 if subtitle else 16)
    paragraph.paragraph_format.space_after = Pt(10 if subtitle else 8)
    run = paragraph.add_run(text)
    format_run(
        run,
        size=12.5 if subtitle else 22,
        bold=not subtitle,
        color=COR_SECUNDARIA if subtitle else COR_PRIMARIA,
    )
    return paragraph


def add_section_title(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(text)
    format_run(run, size=12, bold=True, color=COR_DESTAQUE)
    set_paragraph_border(paragraph)
    return paragraph


def add_subtitle(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(9)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    format_run(run, size=10.8, bold=True, color=COR_PRIMARIA)
    return paragraph


def add_paragraph(document, text, bold_prefix=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.2
    if bold_prefix:
        run = paragraph.add_run(f"{bold_prefix}: ")
        format_run(run, bold=True, color=COR_PRIMARIA)
    run = paragraph.add_run(text)
    format_run(run, color=COR_PRIMARIA)
    return paragraph


def add_bullet_list(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.left_indent = Cm(0.45)
        run = paragraph.add_run(item)
        format_run(run, color=COR_PRIMARIA)


def add_numbered_list(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.left_indent = Cm(0.45)
        run = paragraph.add_run(item)
        format_run(run, color=COR_PRIMARIA)


def add_methodology_steps(document, steps):
    for title, description in steps:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        title_run = paragraph.add_run(f"{title} ")
        format_run(title_run, bold=True, color=COR_PRIMARIA)
        desc_run = paragraph.add_run(f"- {description}")
        format_run(desc_run, color=COR_PRIMARIA)


def add_info_table(document, rows):
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(11.4)

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(4.2)
        cells[1].width = Cm(11.4)
        set_cell_shading(cells[0], COR_CAPA)
        set_cell_shading(cells[1], "FFFFFF")
        for cell in cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p_label = cells[0].paragraphs[0]
        p_value = cells[1].paragraphs[0]
        p_label.paragraph_format.space_after = Pt(0)
        p_value.paragraph_format.space_after = Pt(0)
        r_label = p_label.add_run(label)
        r_value = p_value.add_run(value)
        format_run(r_label, size=9.8, bold=True, color=COR_SECUNDARIA)
        format_run(r_value, color=COR_PRIMARIA)

    document.add_paragraph()


def add_cover_banner(document):
    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eyebrow.paragraph_format.space_before = Pt(18)
    eyebrow.paragraph_format.space_after = Pt(4)
    run = eyebrow.add_run("CONSULTORIA ESPECIALIZADA EM DOCUMENTACAO DE DADOS")
    format_run(run, size=9.2, bold=True, color=COR_MUTED, all_caps=True)

    add_title(document, "PROPOSTA COMERCIAL")
    add_title(document, "Documentacao Tecnica e Executiva para Projetos de Dados", subtitle=True)

    support = document.add_paragraph()
    support.alignment = WD_ALIGN_PARAGRAPH.CENTER
    support.paragraph_format.space_after = Pt(16)
    run = support.add_run(
        "Estrutura comercial pronta para apresentar com clareza tecnica, consistencia executiva e padrao premium de entrega."
    )
    format_run(run, size=10.2, color=COR_SECUNDARIA)


def add_cover_panel(document, config):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(9.8)
    table.columns[1].width = Cm(5.8)

    left_cell, right_cell = table.rows[0].cells
    for cell in (left_cell, right_cell):
        set_cell_border(cell, color="D2DBE4", size="10")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(left_cell, COR_CAPA)
    set_cell_shading(right_cell, "FFFFFF")

    left_title = left_cell.paragraphs[0]
    left_title.paragraph_format.space_after = Pt(6)
    run = left_title.add_run("Escopo da proposta")
    format_run(run, size=10.8, bold=True, color=COR_DESTAQUE)

    left_body = left_cell.add_paragraph()
    left_body.paragraph_format.space_after = Pt(0)
    run = left_body.add_run(
        "Servico de estruturacao documental para transformar artefatos analiticos em materiais tecnicos e executivos claros, organizados e prontos para uso corporativo."
    )
    format_run(run, size=10.2, color=COR_PRIMARIA)

    info_rows = [
        ("Cliente", config["cliente"]),
        ("Projeto", config["projeto"]),
        ("Data", config["data_proposta"]),
        ("Proposta no", config["numero_proposta"]),
    ]
    for idx, (label, value) in enumerate(info_rows):
        paragraph = right_cell.paragraphs[0] if idx == 0 else right_cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        label_run = paragraph.add_run(f"{label}\n")
        format_run(label_run, size=8.8, bold=True, color=COR_MUTED, all_caps=True)
        value_run = paragraph.add_run(value)
        format_run(value_run, size=10.3, bold=label == "Cliente", color=COR_PRIMARIA)

    document.add_paragraph()


def add_pricing_table(document, rows, config):
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(2.2)
    table.columns[1].width = Cm(3.6)
    table.columns[2].width = Cm(8.1)
    table.columns[3].width = Cm(3.5)

    header = table.rows[0].cells
    headers = ["Pacote", "Categoria", "Descricao", "Investimento"]
    for cell, text in zip(header, headers):
        set_cell_shading(cell, COR_TABELA_HEADER)
        set_cell_border(cell, color="243745", size="10")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        format_run(run, size=9.8, bold=True, color=RGBColor(255, 255, 255))

    pacote_selecionado_normalizado = config["pacote_selecionado"].lower()

    for index, (pacote, nome, descricao, investimento) in enumerate(rows):
        cells = table.add_row().cells
        values = [pacote, nome, descricao, investimento]
        highlight = pacote.lower() in pacote_selecionado_normalizado or nome.lower() in pacote_selecionado_normalizado
        for idx, value in enumerate(values):
            cell = cells[idx]
            if highlight:
                set_cell_shading(cell, COR_TABELA_SELECIONADA)
            elif index % 2 == 1:
                set_cell_shading(cell, COR_TABELA_ZEBRA)
            else:
                set_cell_shading(cell, "FFFFFF")
            set_cell_border(cell, color="D7E0E8")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx != 3 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            if idx == 0:
                format_run(run, size=9.9, bold=True, color=COR_DESTAQUE if highlight else COR_PRIMARIA)
            elif idx == 1:
                format_run(run, size=10, bold=True, color=COR_PRIMARIA)
            elif idx == 3:
                format_run(run, size=10, bold=True, color=COR_DESTAQUE if highlight else COR_PRIMARIA)
            else:
                format_run(run, size=9.8, color=COR_PRIMARIA)

    document.add_paragraph()


def add_highlight_block(document, title, items):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(17.0)
    set_cell_shading(cell, COR_BLOCO)
    set_cell_border(cell, color="D0D7E2", size="10")

    title_paragraph = cell.paragraphs[0]
    title_paragraph.paragraph_format.space_after = Pt(5)
    title_run = title_paragraph.add_run(title)
    format_run(title_run, size=11, bold=True, color=COR_DESTAQUE)

    for item in items:
        paragraph = cell.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        format_run(run, color=COR_PRIMARIA)

    document.add_paragraph()


def add_investment_summary(document, config):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(8.0)
    table.columns[1].width = Cm(9.8)

    left_cell, right_cell = table.rows[0].cells
    set_cell_shading(left_cell, COR_TABELA_HEADER)
    set_cell_shading(right_cell, COR_CAPA)
    set_cell_border(left_cell, color="243745", size="10")
    set_cell_border(right_cell, color="D2DBE4", size="10")

    left = left_cell.paragraphs[0]
    left.paragraph_format.space_after = Pt(5)
    run = left.add_run("Pacote recomendado\n")
    format_run(run, size=8.8, bold=True, color=RGBColor(255, 255, 255), all_caps=True)
    run = left.add_run(config["pacote_selecionado"])
    format_run(run, size=12.2, bold=True, color=RGBColor(255, 255, 255))

    right = right_cell.paragraphs[0]
    right.paragraph_format.space_after = Pt(5)
    run = right.add_run("Investimento final proposto\n")
    format_run(run, size=8.8, bold=True, color=COR_MUTED, all_caps=True)
    run = right.add_run(config["investimento_final"])
    format_run(run, size=13.2, bold=True, color=COR_DESTAQUE)

    document.add_paragraph()


def configure_header_footer(section, config):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(f"{config['nome_marca']}  |  Proposta Comercial")
    format_run(run, size=8.8, bold=True, color=COR_SECUNDARIA)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    footer_run = footer.add_run(
        f"{config['seu_nome']}  |  {config['whatsapp']}  |  {config['email']}  |  {config['linkedin']}"
    )
    format_run(footer_run, size=8.5, color=COR_SECUNDARIA)


def add_cover(document, config):
    add_cover_banner(document)
    add_cover_panel(document, config)
    add_investment_summary(document, config)


def build_document(config):
    document = Document()
    configure_page(document.sections[0])
    apply_base_styles(document)
    configure_header_footer(document.sections[0], config)

    add_cover(document, config)
    document.add_section(WD_SECTION.CONTINUOUS)
    content_section = document.sections[-1]
    configure_page(content_section)
    configure_header_footer(content_section, config)

    add_section_title(document, "1. Resumo Executivo")
    for paragraph in RESUMO_EXECUTIVO.split("\n\n"):
        add_paragraph(document, paragraph)

    add_section_title(document, "2. Contexto")
    add_paragraph(document, CONTEXTO_INTRO)
    add_bullet_list(document, CONTEXTO_LISTA)
    add_paragraph(document, CONTEXTO_FINAL)

    add_section_title(document, "3. Objetivo do Trabalho")
    add_bullet_list(document, OBJETIVOS)

    add_section_title(document, "4. Escopo da Entrega")
    add_subtitle(document, "4.1 Documentacao tecnica")
    add_bullet_list(document, ESCOPO_TECNICO)
    add_subtitle(document, "4.2 Material executivo")
    add_bullet_list(document, ESCOPO_EXECUTIVO)
    add_subtitle(document, "4.3 Revisao")
    add_bullet_list(document, ESCOPO_REVISAO)

    add_section_title(document, "5. Entregaveis")
    add_bullet_list(document, ENTREGAVEIS)

    add_section_title(document, "6. Itens Nao Incluidos no Escopo")
    add_paragraph(document, "Para evitar ambiguidades, esta proposta nao inclui, salvo contratacao especifica:")
    add_bullet_list(document, NAO_INCLUIDOS)
    add_paragraph(document, "Caso haja necessidade de atividades adicionais, estas poderao ser avaliadas e orcadas separadamente.")

    add_section_title(document, "7. Metodologia de Trabalho")
    add_methodology_steps(document, METODOLOGIA)

    add_section_title(document, "8. Prazo Estimado")
    add_paragraph(document, "O prazo varia conforme a complexidade do material recebido.")
    add_bullet_list(document, PRAZOS)
    add_paragraph(
        document,
        "O prazo comeca a contar a partir do recebimento integral dos arquivos, alinhamento inicial necessario e confirmacao do aceite da proposta.",
    )

    add_section_title(document, "9. Investimento")
    add_paragraph(
        document,
        "A tabela abaixo apresenta as faixas de contratacao disponiveis. O pacote recomendado e o investimento final proposto aparecem destacados para facilitar avaliacao comercial e tomada de decisao.",
    )
    add_pricing_table(document, PACOTES, config)
    add_investment_summary(document, config)

    add_section_title(document, "10. Condicoes Comerciais")
    add_bullet_list(document, CONDICOES_COMERCIAIS)

    add_section_title(document, "11. Premissas")
    add_bullet_list(document, PREMISSAS)

    add_section_title(document, "12. Diferenciais da Entrega")
    add_highlight_block(document, "Destaques do servico", DIFERENCIAIS)

    add_section_title(document, "13. Proximos Passos")
    add_numbered_list(document, PROXIMOS_PASSOS)

    add_section_title(document, "14. Encerramento")
    add_paragraph(
        document,
        "Fico a disposicao para apoiar a organizacao, valorizacao e comunicacao das entregas do projeto, transformando artefatos tecnicos em materiais mais claros, bem estruturados e prontos para uso profissional.",
    )
    add_info_table(
        document,
        [
            ("Seu nome", config["seu_nome"]),
            ("Nome da marca", config["nome_marca"]),
            ("WhatsApp", config["whatsapp"]),
            ("E-mail", config["email"]),
            ("LinkedIn", config["linkedin"]),
        ],
    )

    return document


def gerar_proposta(
    cliente,
    projeto,
    data_proposta,
    numero_proposta,
    pacote_selecionado,
    investimento_final,
    nome_saida,
):
    """
    Gera uma proposta a partir de parametros dinamicos, preservando o mesmo layout
    usado pelo template padrao do script.
    """
    config = criar_configuracao_proposta(
        cliente=cliente,
        projeto=projeto,
        data_proposta=data_proposta,
        numero_proposta=numero_proposta,
        pacote_selecionado=pacote_selecionado,
        investimento_final=investimento_final,
        nome_saida=nome_saida,
    )
    output_path = Path(__file__).with_name(config["nome_saida"])
    document = build_document(config)
    document.save(output_path)
    return output_path


def main():
    config = criar_configuracao_proposta()
    output_path = gerar_proposta(
        cliente=config["cliente"],
        projeto=config["projeto"],
        data_proposta=config["data_proposta"],
        numero_proposta=config["numero_proposta"],
        pacote_selecionado=config["pacote_selecionado"],
        investimento_final=config["investimento_final"],
        nome_saida=config["nome_saida"],
    )
    print(f"Arquivo gerado com sucesso: {output_path}")


if __name__ == "__main__":
    main()
